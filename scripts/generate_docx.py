#!/usr/bin/env python3
import os, json, argparse, re, glob
import sys
import logging
from docx import Document

# Configure logging
os.makedirs("logs", exist_ok=True)
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[
        logging.FileHandler("logs/pipeline.log"),
        logging.StreamHandler(sys.stderr)
    ]
)
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import imagehash
from PIL import Image

CLOZE_ANSWERS = []

def build_docx_metadata(docx_path):
    abs_path = os.path.abspath(docx_path)
    metadata = {"path": abs_path}
    try:
        stat = os.stat(docx_path)
        metadata["size"] = stat.st_size
        metadata["mtime"] = int(stat.st_mtime)
    except OSError:
        metadata["size"] = None
        metadata["mtime"] = None
    return metadata

def set_shading(properties_element, fill_color):
    try:
        existing = properties_element.find(qn('w:shd'))
        if existing is not None:
            properties_element.remove(existing)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}" w:color="auto" w:val="clear"/>')
        properties_element.append(shading)
    except Exception:
        pass

def is_logo_frame(text):
    if not text:
        return False
    text_lower = text.lower()
    if "gate smashers" in text_lower or "gate smasher" in text_lower:
        words = re.findall(r'\b\w+\b', text_lower)
        if len(words) < 25 or "subscribe" in text_lower or "join" in text_lower or "follow" in text_lower:
            return True
    return False

def are_ocr_texts_similar(text1, text2, threshold=0.95):
    if not text1 or not text2:
        return False
    # Extract unique words with length >= 1 including numbers and variables
    w1 = set(re.findall(r'\b[a-zA-Z0-9_\-\+]{1,}\b', text1.lower()))
    w2 = set(re.findall(r'\b[a-zA-Z0-9_\-\+]{1,}\b', text2.lower()))
    
    if not w1 or not w2:
        return False
        
    common = w1 & w2
    ratio = len(common) / max(len(w1), len(w2))
    return ratio > threshold

def are_rules_similar(r1, r2, threshold=0.85):
    if not r1 or not r2:
        return False
    w1 = set(re.findall(r'\b[a-zA-Z0-9_\-\+]{1,}\b', r1.lower()))
    w2 = set(re.findall(r'\b[a-zA-Z0-9_\-\+]{1,}\b', r2.lower()))
    
    if not w1 or not w2:
        return False
        
    common = w1 & w2
    ratio = len(common) / max(len(w1), len(w2))
    return ratio > threshold


# Helper Functions
def clean_attributions(text):
    if not text or not isinstance(text, str):
        return text
        
    replacements = [
        (r"\bthis is discussed in the lecture\b", "this is discussed"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+explains?\s+that\b", "it is explained that"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+explains?\b", "it is explained"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+introduces?\b", "it is introduced"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+describes?\b", "it is described"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+outlines?\b", "it is outlined"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+demonstrates?\b", "it is demonstrated"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+analyzes?\b", "it is analyzed"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+reviews?\b", "it is reviewed"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+teaches?\b", "it is taught"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+shows?\b", "it is shown"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+discusses?\b", "it is discussed"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+shares?\b", "it is shared"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+mentions?\b", "it is mentioned"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+says?\b", "it is stated"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+warns?\b", "it is warned"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+notes?\b", "it is noted"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+highlights?\b", "it is highlighted"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+compares?\b", "it is compared"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+contrasts?\b", "it is contrasted"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+recommends?\b", "it is recommended"),
        (r"\bthe\s+(?:teacher|lecturer|instructor)\s+points?\s+out\b", "it is pointed out"),
        (r"\blet's look at\b", "examine"),
        (r"\blet's see\b", "examine"),
        (r"\blet's look\b", "examine"),
        (r"\bwe see\b", "is observed"),
        (r"\bwe analyze\b", "analyzing"),
        (r"\bif we analyze\b", "if analyzing"),
    ]
    
    cleaned = text
    for pattern, repl in replacements:
        def make_repl(match):
            val = match.group(0)
            if val and val[0].isupper():
                return repl[0].upper() + repl[1:]
            return repl
        cleaned = re.sub(pattern, make_repl, cleaned, flags=re.IGNORECASE)

    patterns = [
        r'\b(?:the\s+)?(?:teacher|lecturer|instructor|author|speaker|he|she)\s+(?:introduces|outlines|demonstrates|analyzes|shares|reviews|explains|teaches|details|discusses|states|shows|covers|contrasts|warns|says|tells|notes|compares|presents|examines|points\s+out|gives|highlights|identifies|teaches|recommends)\b\s*(?:that|how|the\s+analogy\s+that)?\s*',
        r'\b(?:this\s+)?(?:block|section|chapter|part)\s+(?:covers|discusses|introduces|outlines|explains)\b\s*',
        r'\b(?:he|she|they|we)\s+also\s+(?:reviews|explains|notes|discusses|introduces|outlines|warns|teaches|highlights)\b\s*',
        r'\b(?:finally|first|second|third|then|next|lastly),\s*(?:he|she|they|the\s+teacher|the\s+lecturer|the\s+instructor)\s+(?:notes|says|explains|states|highlights|warns)\b\s*(?:that)?\s*',
    ]
    
    def clean_sentence(sentence):
        s = sentence.strip()
        if not s:
            return s
        
        orig_len = len(s)
        for pattern in patterns:
            prefix_match = re.match(r'^([\s\*\=_<>a-zA-Z]*?)(.*)$', s)
            if prefix_match:
                prefix, body = prefix_match.group(1), prefix_match.group(2)
                new_body, count = re.subn(r'^' + pattern, '', body, flags=re.IGNORECASE)
                if count > 0:
                    new_body = re.sub(r'^(?:\s*[:\-–—])?\s*', '', new_body)
                    s = prefix + new_body
        
        if len(s) < orig_len:
            for idx, char in enumerate(s):
                if char.isalpha():
                    s = s[:idx] + s[idx].upper() + s[idx+1:]
                    break
        return s

    lines = cleaned.split('\n')
    cleaned_lines = []
    for line in lines:
        parts = re.split(r'([.!?]\s+)', line)
        cleaned_parts = []
        for part in parts:
            if re.match(r'^[.!?]\s+$', part):
                cleaned_parts.append(part)
            else:
                cleaned_parts.append(clean_sentence(part))
        cleaned_lines.append("".join(cleaned_parts))
    
    return "\n".join(cleaned_lines)

def find_matching_brace(text, open_idx):
    count = 0
    for i in range(open_idx, len(text)):
        if text[i] == '{':
            count += 1
        elif text[i] == '}':
            count -= 1
            if count == 0:
                return i
    return -1

def parse_latex_fractions(text):
    while r'\frac' in text:
        idx = text.find(r'\frac')
        open1 = text.find('{', idx)
        if open1 == -1:
            break
        close1 = find_matching_brace(text, open1)
        if close1 == -1:
            break
        open2 = text.find('{', close1)
        if open2 == -1 or open2 > close1 + 2:
            break
        close2 = find_matching_brace(text, open2)
        if close2 == -1:
            break
        
        num = text[open1+1:close1]
        den = text[open2+1:close2]
        
        num = parse_latex_fractions(num)
        den = parse_latex_fractions(den)
        
        text = text[:idx] + f"({num}) / ({den})" + text[close2+1:]
    return text

def format_math_text(text, is_code=False):
    if not text:
        return text
    if is_code:
        return text
        
    # Remove standard math delimiters first
    text = text.replace(r'\(', '').replace(r'\)', '').replace(r'\[', '').replace(r'\]', '')
    text = text.replace('$$', '')
    text = re.sub(r'(?<!\\)\$', '', text)
    
    # Strip \text{} commands
    text = re.sub(r'\\text\s*\{([^}]*)\}', r'\1', text)
    
    # Parse fractions recursively
    while True:
        match = re.search(r'\\frac\s*\{([^}]*)\}\s*\{([^}]*)\}', text)
        if not match:
            break
        text = text.replace(match.group(0), f"({match.group(1)})/({match.group(2)})")
        
    # Normalize square roots recursively
    while True:
        match = re.search(r'\\sqrt\s*\{([^}]*)\}', text)
        if not match:
            break
        text = text.replace(match.group(0), f"√({match.group(1)})")
    text = re.sub(r'\\sqrt\s*([a-zA-Z0-9])', r'√\1', text)
    text = re.sub(r'\bsqrt\((.*?)\)', r'√(\1)', text)
    
    # Convert nPr and nCr notation
    superscripts = {
        '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
        '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹',
        '-': '⁻', '+': '⁺', '=': '⁼', '(': '⁽', ')': '⁾',
        'a': 'ᵃ', 'b': 'ᵇ', 'c': 'ᶜ', 'd': 'ᵈ', 'e': 'ᵉ', 'f': 'ᶠ', 'g': 'ᵍ', 'h': 'ʰ', 'i': 'ⁱ', 'j': 'ʲ', 'k': 'ᵏ', 'l': 'ˡ', 'm': 'ᵐ', 'n': 'ⁿ', 'o': 'ᵒ', 'p': 'ᵖ', 'r': 'ʳ', 's': 'ˢ', 't': 'ᵗ', 'u': 'ᵘ', 'v': 'ᵛ', 'w': 'ʷ', 'x': 'ˣ', 'y': 'ʸ', 'z': 'ᶻ',
        'A': 'ᴬ', 'B': 'ᴮ', 'D': 'ᴰ', 'E': 'ᴱ', 'G': 'ᴳ', 'H': 'ᴴ', 'I': 'ᴵ', 'J': 'ᴶ', 'K': 'ᴷ', 'L': 'ᴸ', 'M': 'ᴹ', 'N': 'ᴺ', 'O': 'ᴼ', 'P': 'ᴾ', 'R': 'ᴿ', 'T': 'ᵀ', 'U': 'ᵁ', 'V': 'ⱽ', 'W': 'ᵂ'
    }
    subscripts = {
        '0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄',
        '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉',
        '-': '₋', '+': '₊', '=': '₌', '(': '₍', ')': '₎',
        'a': 'ₐ', 'e': 'ₑ', 'h': 'ₕ', 'i': 'ᵢ', 'j': 'ⱼ', 'k': 'ₖ', 'l': 'ₗ', 'm': 'ₘ', 'n': 'ₙ', 'o': 'ₒ', 'p': 'ₚ', 'r': 'ᵣ', 's': 'ₛ', 't': 'ₜ', 'u': 'ᵤ', 'v': 'ᵥ', 'x': 'ₓ',
        'A': 'ₐ', 'B': 'B', 'D': 'D', 'E': 'ₑ', 'G': 'G', 'H': 'ₕ', 'I': 'ᵢ', 'J': 'ⱼ', 'K': 'ₖ', 'L': 'ₗ', 'M': 'ₘ', 'N': 'ₙ', 'O': 'ₒ', 'P': 'ₚ', 'R': 'ᵣ', 'T': 'ₜ', 'U': 'ᵤ', 'V': 'ᵥ', 'W': 'W'
    }
    
    def repl_npr_ncr(m):
        n_val = m.group(1)
        type_val = m.group(2)
        r_val = m.group(3)
        n_sup = "".join(superscripts.get(c, c) for c in n_val)
        r_sub = "".join(subscripts.get(c, c) for c in r_val)
        return f"{n_sup}{type_val}{r_sub}"
        
    text = re.sub(r'\^\{?([0-9a-zA-Z]+)\}?\\text\{([PC])\}_\{?([0-9a-zA-Z]+)\}?', repl_npr_ncr, text)
    text = re.sub(r'\^\{?([0-9a-zA-Z]+)\}?([PC])_\{?([0-9a-zA-Z]+)\}?', repl_npr_ncr, text)

    # Standard LaTeX math symbols replacements
    replacements = {
        r'\\times': '×',
        r'\\cdot': '·',
        r'\\approx': '≈',
        r'\\leq': '≤',
        r'\\le\b': '≤',
        r'\\geq': '≥',
        r'\\ge\b': '≥',
        r'\\neq': '≠',
        r'\\ne\b': '≠',
        r'\\sigma': 'σ',
        r'\\alpha': 'α',
        r'\\beta': 'β',
        r'\\delta': 'δ',
        r'\\theta': 'θ',
        r'\\pi': 'π',
        r'\\mu': 'μ',
        r'\\lambda': 'λ',
        r'\\omega': 'ω',
        r'\\phi': 'φ',
        r'\\oplus': '⊕',
        r'\\sum': '∑',
        r'\\infty': '∞',
        r'\\cup': '∪',
        r'\\cap': '∩',
        r'\\in\b': '∈',
        r'\\notin\b': '∉',
        r'\\subset\b': '⊂',
        r'\\subseteq\b': '⊆',
        r'\\to\b': '→',
        r'\\rightarrow\b': '→',
        r'\\gets\b': '←',
        r'\\leftarrow\b': '←',
        r'\\iff\b': '⇔',
        r'\\implies\b': '⇒',
        r'\\pm': '±',
        r'\\dots': '...',
        r'\\lfloor': '⌊',
        r'\\rfloor': '⌋',
        r'\\langle': '⟨',
        r'\\rangle': '⟩',
        r'\\div': '÷',
        r'\\prod': '∏',
        r'\\int': '∫',
        r'\\forall': '∀',
        r'\\exists': '∃',
        r'\\emptyset': '∅',
        r'\\triangle': '△',
        r'\\angle': '∠',
        r'\\degree': '°',
        r'\\circ': '°',
        r'\\log': 'log',
        r'\\ln': 'ln',
        r'\\sin': 'sin',
        r'\\cos': 'cos',
        r'\\tan': 'tan',
    }
    
    for pattern, rep in replacements.items():
        text = re.sub(pattern, rep, text)

    # Support pmod formatting
    text = re.sub(r'\\pmod\s*\{([^}]*)\}', r'(mod \1)', text)
    text = re.sub(r'\\pmod\s*([a-zA-Z0-9])', r'(mod \1)', text)
    
    # ASCII equivalents for multiplication
    text = re.sub(r'(?=[0-9a-zA-Z\)])\s*\*\s*(?=[0-9a-zA-Z\(√])', ' × ', text)
    text = re.sub(r'(?=[0-9a-zA-Z\)\]])\s+\*\s+(?=[0-9a-zA-Z\(√\[])', ' × ', text)
    
    # Greek words fallback mapping (for un-backslashed or backslashed Greek words)
    greek = {
        'alpha': 'α', 'beta': 'β', 'gamma': 'γ', 'delta': 'δ', 'epsilon': 'ε',
        'zeta': 'ζ', 'eta': 'η', 'theta': 'θ', 'iota': 'ι', 'kappa': 'κ',
        'lambda': 'λ', 'mu': 'μ', 'nu': 'ν', 'xi': 'ξ', 'pi': 'π', 'rho': 'ρ',
        'sigma': 'σ', 'tau': 'τ', 'upsilon': 'υ', 'phi': 'φ', 'chi': 'χ',
        'psi': 'ψ', 'omega': 'ω'
    }
    for g_str, g_uni in greek.items():
        text = re.sub(rf'\\?\b{g_str}\b', g_uni, text)
        text = re.sub(rf'\\?\b{g_str.capitalize()}\b', g_uni, text)
        
    # Clean stray LaTeX backslashes and leftover $ after all conversions
    text = re.sub(r'\\(?=[a-zA-Z])', '', text)  # stray \letter
    text = re.sub(r'(?<!\\)\$', '', text)  # stray $ not escaped

    return text

def add_rich_runs(p, text, default_bold=False, default_italic=False, default_underline=False, default_color_rgb=None, default_highlight=None, default_font_size=None):
    if not text:
        return
        
    if isinstance(text, dict):
        text = text.get('text', text.get('sentence', str(text)))
    elif isinstance(text, list):
        text = ' '.join(str(x) for x in text)
    else:
        text = str(text)

    # Pre-convert LaTeX math commands to Unicode equivalents at the very beginning
    text = text.replace(r'\pm', '±')
    text = text.replace(r'\times', '×')
    text = text.replace(r'\leq', '≤')
    text = text.replace(r'\geq', '≥')
    text = text.replace(r'\neq', '≠')
    text = text.replace(r'\approx', '≈')
    text = text.replace(r'\div', '÷')
    text = text.replace(r'\cdot', '·')
    text = text.replace(r'\rightarrow', '→')
    text = text.replace(r'\leftarrow', '←')
    text = text.replace(r'\Rightarrow', '⇒')
    text = text.replace(r'\infty', '∞')
    text = text.replace(r'\min', 'min')
    text = text.replace(r'\max', 'max')
    text = text.replace(r'\sum', '∑')
    text = text.replace(r'\prod', '∏')
    text = text.replace(r'\int', '∫')
    text = text.replace(r'\forall', '∀')
    text = text.replace(r'\exists', '∃')
    text = re.sub(r'(?<![a-zA-Z])\\in(?![a-zA-Z])', '∈', text)
    text = text.replace(r'\notin', '∉')
    text = text.replace(r'\subset', '⊂')
    text = text.replace(r'\supset', '⊃')
    text = text.replace(r'\cup', '∪')
    text = text.replace(r'\cap', '∩')
    text = text.replace(r'\emptyset', '∅')
    text = text.replace(r'\triangle', '△')
    text = text.replace(r'\angle', '∠')
    text = text.replace(r'\degree', '°')
    text = text.replace(r'\circ', '°')
    text = text.replace(r'\log', 'log')
    text = text.replace(r'\ln', 'ln')
    text = text.replace(r'\sin', 'sin')
    text = text.replace(r'\cos', 'cos')
    text = text.replace(r'\tan', 'tan')

    # Clean attributions
    text = clean_attributions(text)
    
    # Normalize potential LLM-generated code tags to standard no-underscore tags
    text = text.replace('<code_inline>', '<codeinline>').replace('</code_inline>', '</codeinline>')
    text = text.replace('<code>', '<codeinline>').replace('</code>', '</codeinline>')
    text = text.replace('<code_block>', '<codeblock>').replace('</code_block>', '</codeblock>')

    # Convert code blocks and inline code backticks to custom XML tags
    # We use no-underscore tag names 'codeblock' and 'codeinline' to prevent 
    # the subscript parser _([a-zA-Z0-9]) below from matching and splitting the tags.
    text = re.sub(r'```(.*?)```', r'<codeblock>\1</codeblock>', text, flags=re.DOTALL)
    text = re.sub(r'`([^`]+)`', r'<codeinline>\1</codeinline>', text)
    # Remove any unmatched backticks
    text = text.replace('`', '')
    
    text = re.sub(r'([a-zA-Z0-9])_?\{?\\text\{([^}]+)\}\}?', r'\1<sub>\2</sub>', text)
    text = re.sub(r'\\text\{([^}]+)\}', r'\1', text)

    # Pre-convert LaTeX exponents and subscripts to HTML-like tags
    text = re.sub(r'\^\{([^}]+)\}', r'<sup>\1</sup>', text)
    text = re.sub(r'_\{([^}]+)\}', r'<sub>\1</sub>', text)
    text = re.sub(r'\^([a-zA-Z0-9])', r'<sup>\1</sup>', text)
    text = re.sub(r'(?<=[a-zA-Z0-9])_([a-zA-Z0-9])', r'<sub>\1</sub>', text)
    # Pre-convert markdown bold, highlight, and italic to HTML-like tags
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'==(.*?)==', r'<highlight color="BLUE">\1</highlight>', text)
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    text = re.sub(r'(?<!\w)_(.*?)_(?!\w)', r'<i>\1</i>', text)
    
    # Parse HTML-like tags: <b>, <i>, <u>, <color rgb="...">, <highlight color="...">, <cloze>, <sup>, <sub>
    pattern = re.compile(r'(</?[a-zA-Z_]+(?:\s+[a-zA-Z_]+=(?:["\'][^"\']*["\']|[^>\s]+))*\s*/?>)')
    tokens = pattern.split(text)
    
    style_stack = []
    
    for token in tokens:
        if not token:
            continue
        if token.startswith('<') and token.endswith('>'):
            is_closing = token.startswith('</')
            tag_name_match = re.match(r'</?([a-zA-Z_]+)', token)
            if tag_name_match:
                tag_name = tag_name_match.group(1).lower()
                if is_closing:
                    for i in range(len(style_stack) - 1, -1, -1):
                        if style_stack[i]['tag'] == tag_name:
                            style_stack.pop(i)
                            break
                else:
                    attrs = {}
                    attr_matches = re.findall(r'([a-zA-Z_]+)=(?:["\']([^"\']*)["\']|([a-zA-Z0-9_#]+))', token)
                    for k, v1, v2 in attr_matches:
                        v = v1 if v1 else v2
                        attrs[k.lower()] = v
                    style_stack.append({'tag': tag_name, 'attrs': attrs})
        else:
            is_code = any(style['tag'] in ('codeblock', 'codeinline', 'code_block', 'code_inline') for style in style_stack)
            cleaned_text = format_math_text(token, is_code=is_code)
            run = p.add_run(cleaned_text)
            run.font.name = 'Consolas' if is_code else 'Calibri'
            if default_font_size is not None:
                run.font.size = default_font_size
            
            bold = default_bold
            italic = default_italic
            underline = default_underline
            color_rgb = default_color_rgb
            highlight_color = default_highlight
            superscript = False
            subscript = False
            
            for style in style_stack:
                t = style['tag']
                attrs = style['attrs']
                if t == 'b':
                    bold = True
                elif t == 'i':
                    italic = True
                elif t == 'u':
                    underline = True
                elif t == 'color':
                    color_rgb = attrs.get('rgb')
                elif t == 'highlight':
                    highlight_color = attrs.get('color')
                elif t == 'sup':
                    superscript = True
                elif t == 'sub':
                    subscript = True
                elif t == 'cloze':
                    underline = True
                    color_rgb = RGBColor(0x00, 0x70, 0xC0)
            
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if underline:
                run.underline = True
            if superscript:
                run.font.superscript = True
            if subscript:
                run.font.subscript = True
            if color_rgb:
                try:
                    if isinstance(color_rgb, RGBColor):
                        run.font.color.rgb = color_rgb
                    else:
                        hex_color = color_rgb.lstrip('#')
                        run.font.color.rgb = RGBColor(
                            int(hex_color[0:2], 16),
                            int(hex_color[2:4], 16),
                            int(hex_color[4:6], 16)
                        )
                except Exception as e:
                    pass
            if highlight_color or is_code:
                if is_code:
                    hex_fill = 'EAECEE' # Light gray shading for code blocks/inline code
                else:
                    color_map = {
                        'YELLOW': 'F5CBA7', 
                        'GREEN': 'A3E4D7',  
                        'BLUE': 'AED6F1',   
                        'GRAY': 'D5D8DC',   
                        'RED': 'F5B7B1',    
                        'ORANGE': 'F5CBA7', 
                        'PURPLE': 'D2B4DE', 
                    }
                    hc_upper = highlight_color.upper()
                    hex_fill = color_map.get(hc_upper, 'AED6F1')
                try:
                    rPr = run._r.get_or_add_rPr()
                    set_shading(rPr, hex_fill)
                except Exception as e:
                    pass

            cloze_style = next((s for s in style_stack if s['tag'] == 'cloze'), None)
            if cloze_style:
                if 'cloze_id' not in cloze_style:
                    attrs = cloze_style['attrs']
                    answer = attrs.get('answer', '')
                    hint = attrs.get('hint', '')
                    
                    global CLOZE_ANSWERS
                    cloze_id = len(CLOZE_ANSWERS) + 1
                    CLOZE_ANSWERS.append({
                        'id': cloze_id,
                        'answer': answer,
                        'hint': hint
                    })
                    cloze_style['cloze_id'] = cloze_id
                    
                    # Add superscript run showing the cloze ID
                    super_run = p.add_run(f"{cloze_id}")
                    super_run.font.superscript = True
                    super_run.font.name = 'Calibri'
                    if default_font_size is not None:
                        super_run.font.size = max(Pt(6), default_font_size - Pt(2))
                    else:
                        super_run.font.size = Pt(8.5)
                    super_run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

def add_custom_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    # Do NOT set explicit dark colors on headings so they are compatible with dark mode auto-inversion
    for r in h.runs:
        r.font.name = 'Calibri'
    return h

def add_bold_prefix(p, text):
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    return run

def add_revision_box(doc, bullets, rule=None):
    parts = ['[⚡ Quick Rev]'] + [f'• {b}' for b in bullets]
    if rule:
        parts.append(f'Rule: {rule}')
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    for idx, part in enumerate(parts):
        if idx > 0:
            p.add_run('\n')
        if part.startswith('[⚡ Quick Rev]'):
            add_rich_runs(p, part, default_bold=True)
        else:
            add_rich_runs(p, part)
            
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D6EAF8" w:val="clear"/>')
    p._p.get_or_add_pPr().append(shading)
    
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
    return p

def add_trap(doc, text):
    p = doc.add_paragraph()
    p.add_run("🚨 TRAP: ").bold = True
    add_rich_runs(p, text, default_bold=True)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    return p

def add_trick(doc, text):
    p = doc.add_paragraph()
    p.add_run("💡 TRICK: ").bold = True
    add_rich_runs(p, text, default_bold=True)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    return p

def add_quote(doc, text):
    p = doc.add_paragraph()
    p.add_run("📝 QUOTE: ").bold = True
    add_rich_runs(p, f'"{text}"', default_italic=True)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    return p

def add_correction(doc, wrong, correct):
    p = doc.add_paragraph()
    p.add_run("Wrong: ").bold = True
    run_wrong = p.add_run(format_math_text(wrong))
    run_wrong.font.strike = True
    run_wrong.font.name = 'Calibri'
    p.add_run(' → ')
    p.add_run("Correct: ").bold = True
    run_correct = p.add_run(format_math_text(correct))
    run_correct.bold = True
    run_correct.font.name = 'Calibri'
    return p

def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.text = ""
        add_rich_runs(p, header, default_bold=True, default_color_rgb="FFFFFF")
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2C3E50" w:val="clear"/>')  # Keep dark header
        cell._element.get_or_add_tcPr().append(shading)
    for ri, row_data in enumerate(rows):
        for ci, value in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            p = cell.paragraphs[0]
            p.text = ""
            add_rich_runs(p, str(value))
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
            if ri % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F3F4" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shading)
                for run in p.runs:
                    if run.font.color.rgb is None:
                        run.font.color.rgb = RGBColor(0, 0, 0)
    return table

def add_image_if_exists(doc, path, width=Inches(4.5)):
    if os.path.exists(path):
        try:
            doc.add_picture(path, width=width)
            return True
        except Exception as e:
            logger.warning(f"Failed to insert image {path}: {e}")
            p = doc.add_paragraph(f"[Image failed to load: {os.path.basename(path)}]")
            p.italic = True
            return False
    return False

def add_shaded_formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_rich_runs(p, text, default_bold=True, default_font_size=Pt(11))
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EAECEE" w:val="clear"/>')
    p._p.get_or_add_pPr().append(shading)
    for run in p.runs:
        run.font.name = 'Calibri'
        if run.font.color.rgb is None:
            run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_super_key_additions(doc):
    # Addition 1: Master Formula Clarification
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(8)
    p1.paragraph_format.space_after = Pt(8)
    p1.paragraph_format.left_indent = Inches(0.25)
    p1.paragraph_format.right_indent = Inches(0.25)
    
    r = p1.add_run("💡 Master Formula (General Case):\n\n")
    r.bold = True
    
    r = p1.add_run("Number of Super Keys = 2")
    r.bold = True
    r_sup = p1.add_run("n − k")
    r_sup.bold = True
    r_sup.font.superscript = True
    r = p1.add_run("\n\n")
    
    r = p1.add_run("where:\n")
    
    r = p1.add_run("• n ")
    r.bold = True
    r = p1.add_run("= Total number of attributes in Relation ")
    r = p1.add_run("R\n")
    r.bold = True
    
    r = p1.add_run("• k ")
    r.bold = True
    r = p1.add_run("= Number of attributes in the ")
    r = p1.add_run("Candidate Key\n\n")
    r.bold = True
    
    r = p1.add_run("Special Case (Lecture Shortcut):\n")
    r.bold = True
    
    r = p1.add_run("If the Candidate Key contains only ")
    r = p1.add_run("one attribute (k = 1)")
    r.bold = True
    r = p1.add_run(", then the formula becomes:\n\n")
    
    r = p1.add_run("Number of Super Keys = 2")
    r.bold = True
    r_sup2 = p1.add_run("n − 1")
    r_sup2.bold = True
    r_sup2.font.superscript = True
    r = p1.add_run("\n\n")
    
    r = p1.add_run("This is the shortcut used in the lecture. It is simply a special case of the general formula.\n\n")
    
    r = p1.add_run("Note: ")
    r.bold = True
    r = p1.add_run("Some sources may display the formula as ")
    r = p1.add_run("2n-k")
    r.bold = True
    r_sup3 = p1.add_run("n−k")
    r_sup3.bold = True
    r_sup3.font.superscript = True
    r = p1.add_run(" (2 raised to the power of ")
    r = p1.add_run("(n−k)")
    r.italic = True
    r = p1.add_run("), ")
    r = p1.add_run("not")
    r.bold = True
    r = p1.add_run(" (2 × n − k).\n\n")
    
    r = p1.add_run("Reference: ")
    r.bold = True
    r = p1.add_run("GeeksforGeeks – ")
    r = p1.add_run("Number of Possible Super Keys in DBMS")
    r.italic = True
    r = p1.add_run(".")
    
    shading1 = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4F8" w:val="clear"/>')
    p1._p.get_or_add_pPr().append(shading1)
    for run in p1.runs:
        run.font.name = 'Calibri'
        if run.font.size is None:
            run.font.size = Pt(10)
        if run.font.color.rgb is None:
            run.font.color.rgb = RGBColor(0, 0, 0)
        
    doc.add_paragraph()

    # Addition 2: Multiple Candidate Keys Exam Note
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(8)
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.left_indent = Inches(0.25)
    p2.paragraph_format.right_indent = Inches(0.25)
    
    r = p2.add_run("⚠️ Exam Note (Multiple Candidate Keys):\n\n")
    r.bold = True
    
    r = p2.add_run("The formula ")
    r = p2.add_run("2")
    r.bold = True
    r_sup4 = p2.add_run("n−k")
    r_sup4.bold = True
    r_sup4.font.superscript = True
    r = p2.add_run(" directly applies when counting super keys generated from ")
    r = p2.add_run("a single candidate key")
    r.bold = True
    r = p2.add_run(".\n\n")
    
    r = p2.add_run("If a relation contains ")
    r = p2.add_run("multiple candidate keys")
    r.bold = True
    r = p2.add_run(", you ")
    r = p2.add_run("cannot simply add")
    r.bold = True
    r = p2.add_run(" the number of super keys obtained from each candidate key, because some super keys will contain more than one candidate key and would therefore be counted multiple times.\n\n")
    
    r = p2.add_run("To obtain the correct answer, use the ")
    r = p2.add_run("Inclusion–Exclusion Principle")
    r.bold = True
    r = p2.add_run(":\n\n")
    
    r = p2.add_run("Total Super Keys\n")
    r.bold = True
    r.font.size = Pt(11)
    
    r = p2.add_run("  (Super Keys from Candidate Key 1)\n")
    r = p2.add_run("+ (Super Keys from Candidate Key 2)\n")
    r = p2.add_run("− (Common Super Keys containing both Candidate Keys)\n\n")
    
    r = p2.add_run("For three or more candidate keys, extend the Inclusion–Exclusion Principle accordingly.\n\n")
    
    r = p2.add_run("This prevents ")
    r = p2.add_run("double-counting")
    r.bold = True
    r = p2.add_run(" of common super keys.\n\n")
    
    r = p2.add_run("Reference: ")
    r.bold = True
    r = p2.add_run("GeeksforGeeks – ")
    r = p2.add_run("Number of Possible Super Keys in DBMS")
    r.italic = True
    r = p2.add_run(" (Examples 4–9).")
    
    shading2 = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4F8" w:val="clear"/>')
    p2._p.get_or_add_pPr().append(shading2)
    for run in p2.runs:
        run.font.name = 'Calibri'
        if run.font.size is None:
            run.font.size = Pt(10)
        if run.font.color.rgb is None:
            run.font.color.rgb = RGBColor(0, 0, 0)
        
    doc.add_paragraph()

def add_cornell_block(doc, cue_text, content, srs_tag=None):
    table = doc.add_table(rows=1, cols=2)
    
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="nil"/>'
        '  <w:left w:val="nil"/>'
        '  <w:bottom w:val="nil"/>'
        '  <w:right w:val="nil"/>'
        '  <w:insideH w:val="nil"/>'
        '  <w:insideV w:val="nil"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    
    look = tblPr.find(qn('w:tblLook'))
    if look is not None:
        look.set(qn('w:firstRow'), '0')
        look.set(qn('w:lastRow'), '0')
        look.set(qn('w:firstColumn'), '0')
        look.set(qn('w:lastColumn'), '0')
        
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = parse_xml(f'<w:tblLayout {nsdecls("w")}/>')
        tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = parse_xml(f'<w:tblGrid {nsdecls("w")}/>')
        tbl.insert(1, tblGrid)
    
    for col in tblGrid.findall(qn('w:gridCol')):
        tblGrid.remove(col)
        
    for w_val in ["2880", "6480"]:
        col = parse_xml(f'<w:gridCol {nsdecls("w")} w:w="{w_val}"/>')
        tblGrid.append(col)

    cell_cue = table.cell(0, 0)
    cell_note = table.cell(0, 1)
    
    cell_cue.width = Inches(2.0)
    cell_note.width = Inches(4.5)
    
    for cell in [cell_cue, cell_note]:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            '  <w:top w:w="120" w:type="dxa"/>'
            '  <w:left w:w="120" w:type="dxa"/>'
            '  <w:bottom w:w="120" w:type="dxa"/>'
            '  <w:right w:w="120" w:type="dxa"/>'
            '</w:tcMar>'
        )
        tcPr.append(tcMar)
        
        cell_borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tcBorders>'
        )
        tcPr.append(cell_borders)
        
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D5DBDB" w:val="clear"/>')
    cell_cue._tc.get_or_add_tcPr().append(shading)
    
    p_cue = cell_cue.paragraphs[0]
    p_cue.paragraph_format.space_before = Pt(0)
    p_cue.paragraph_format.space_after = Pt(4)
    cue_text_clean = clean_attributions(cue_text).replace('`', '')
    run_cue = p_cue.add_run(cue_text_clean)
    run_cue.font.name = 'Calibri'
    run_cue.font.size = Pt(11)
    run_cue.italic = True
    run_cue.font.color.rgb = RGBColor(0, 0, 0)
    
    if srs_tag:
        p_srs = cell_cue.add_paragraph()
        p_srs.paragraph_format.space_before = Pt(4)
        p_srs.paragraph_format.space_after = Pt(0)
        display_srs = srs_tag.replace("SRS:", "Review Interval:")
        run_srs = p_srs.add_run(f"⏳ {display_srs}")
        run_srs.font.name = 'Calibri'
        run_srs.font.size = Pt(9.5)
        run_srs.bold = True
        run_srs.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
        
    p_note = cell_note.paragraphs[0]
    p_note.paragraph_format.space_before = Pt(0)
    p_note.paragraph_format.space_after = Pt(4)
    
    if isinstance(content, list):
        first = True
        for item in content:
            val = item.get('value', '') if isinstance(item, dict) else str(item)
            lines = val.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                p_new = p_note if first else cell_note.add_paragraph()
                first = False
                p_new.paragraph_format.space_before = Pt(0)
                p_new.paragraph_format.space_after = Pt(4)
                add_rich_runs(p_new, line, default_font_size=Pt(11))
                
                # Apply shading based on prefix for Gate 22 and visual conformity
                line_clean = line.replace("<b>", "").replace("</b>", "").strip()
                if line_clean.startswith("💡 Student Note / Doubt") or line_clean.startswith("📝 Student Note:"):
                    try:
                        shading_el = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4F8" w:val="clear"/>')
                        p_new._p.get_or_add_pPr().append(shading_el)
                        for r in p_new.runs:
                            if r.font.color.rgb is None:
                                r.font.color.rgb = RGBColor(0, 0, 0)
                    except Exception:
                        pass
                elif line_clean.startswith("💡 Teacher's Intuition & Analogies:"):
                    try:
                        shading_el = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF2CC" w:val="clear"/>')
                        p_new._p.get_or_add_pPr().append(shading_el)
                        for r in p_new.runs:
                            if r.font.color.rgb is None:
                                r.font.color.rgb = RGBColor(0, 0, 0)
                    except Exception:
                        pass
    else:
        val = str(content)
        lines = val.split('\n')
        first = True
        for line in lines:
            line = line.strip()
            if not line:
                continue
            p_new = p_note if first else cell_note.add_paragraph()
            first = False
            p_new.paragraph_format.space_before = Pt(0)
            p_new.paragraph_format.space_after = Pt(4)
            add_rich_runs(p_new, line, default_font_size=Pt(11))
            
            # Apply shading based on prefix for Gate 22 and visual conformity
            line_clean = line.replace("<b>", "").replace("</b>", "").strip()
            if line_clean.startswith("💡 Student Note / Doubt") or line_clean.startswith("📝 Student Note:"):
                try:
                    shading_el = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4F8" w:val="clear"/>')
                    p_new._p.get_or_add_pPr().append(shading_el)
                    for r in p_new.runs:
                        if r.font.color.rgb is None:
                            r.font.color.rgb = RGBColor(0, 0, 0)
                except Exception:
                    pass
            elif line_clean.startswith("💡 Teacher's Intuition & Analogies:"):
                try:
                    shading_el = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF2CC" w:val="clear"/>')
                    p_new._p.get_or_add_pPr().append(shading_el)
                    for r in p_new.runs:
                        if r.font.color.rgb is None:
                            r.font.color.rgb = RGBColor(0, 0, 0)
                except Exception:
                    pass
        
    doc.add_paragraph()

def add_formatted_explanation_paragraphs(doc, text):
    if not text:
        return
        
    if not isinstance(text, str):
        text = str(text)
        
    # Clean attributions first
    text = clean_attributions(text)
        
    # Pre-convert markdown bold, highlight, and italic to HTML-like tags
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'==(.*?)==', r'<highlight color="BLUE">\1</highlight>', text)
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    text = re.sub(r'(?<!\w)_(.*?)_(?!\w)', r'<i>\1</i>', text)
    
    # Split text into lines first
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if line starts with list/bullet markers
        if line.startswith(('•', '-', '*')):
            sent_clean = line.lstrip('•-* ').strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(6)
            add_rich_runs(p, sent_clean)
        elif re.match(r'^\d+[\.\)]', line):
            match = re.match(r'^(\d+[\.\)])\s*(.*)', line)
            prefix = match.group(1)
            body = match.group(2)
            # Use normal paragraph with left indent to avoid Word's automatic list numbering continuation bug
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(6)
            p.add_run(prefix + " ").bold = True
            add_rich_runs(p, body)
        else:
            # Render the line, checking if it exceeds 350 characters
            if len(line) > 200:
                sentences = re.split(r'(?<!\bi\.e)(?<!\be\.g)(?<!\bvs)(?<!\bapprox)(?<!\bfig)(?<!\bno)(?<=[.!?])\s+(?=[A-Z"“0-9])', line)
                chunks = []
                for i in range(0, len(sentences), 3):
                    chunk = sentences[i:i+3]
                    chunk_text = " ".join(chunk).strip()
                    if chunk_text:
                        chunks.append(chunk_text)
                for chunk_text in chunks:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(6)
                    add_rich_runs(p, chunk_text)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                add_rich_runs(p, line)

def get_explicit_answer(ex, working_text):
    answer = ex.get('answer')
    if answer:
        return answer

    if not working_text:
        return ""

    # Prefer the final stated conclusion when the author did not supply an explicit answer.
    for line in reversed([ln.strip() for ln in working_text.split('\n') if ln.strip()]):
        if line.lower().startswith('therefore'):
            return line.split('Therefore,', 1)[-1].strip().lstrip(':').strip()

    # Fallback only if the working text explicitly ends with a result-like line.
    if '->' in working_text:
        return working_text.split('->')[-1].strip()

    return ""


def find_best_embedded_screenshot(context_text, embedded_screenshots):
    if not embedded_screenshots or not context_text:
        return None
    
    best_img = None
    best_score = 0.0
    
    context_lower = context_text.lower()
    # Extract lowercase alphanumeric words of length >= 2 including numbers and variables
    context_words = set(re.findall(r'\b[a-zA-Z0-9_\-\+]{2,}\b', context_lower))
    
    for ss in embedded_screenshots:
        ocr_lower = ss.get('ocr_text', '').lower()
        if not ocr_lower:
            continue
        ocr_words = set(re.findall(r'\b[a-zA-Z0-9_\-\+]{2,}\b', ocr_lower))
        if not ocr_words:
            continue
            
        intersection = context_words & ocr_words
        overlap_ocr = len(intersection) / len(ocr_words)
        
        union = len(context_words | ocr_words)
        jaccard = len(intersection) / union if union > 0 else 0
        
        # Robust scoring: prioritize ocr overlap and Jaccard similarity, with a small bonus for absolute match count
        score = overlap_ocr * 0.4 + jaccard * 0.6 + len(intersection) * 0.02
        
        if score > best_score:
            best_score = score
            best_img = ss.get('image_path')
            
    if best_score > 0.30:
        return best_img
    return None

def get_vm_image_path(vm, block_id, slides, timestamp_to_frame, frames, embedded_screenshots=None, context_text=""):
    ts = vm.get('timestamp', '').rstrip('*')
    v_type = vm.get('type', 'board')
    
    img_path = None
    frame_fname = None
    if v_type == 'slide':
        slide_num = vm.get('slide_number')
        if slide_num and slides:
            is_ref = vm.get('source') == 'reference'
            for s in slides:
                if s.get('slide_number') == slide_num:
                    s_path = s.get('image_path', '')
                    if is_ref and 'reference_pages' not in s_path:
                        continue
                    if not is_ref and 'reference_pages' in s_path:
                        continue
                    img_path = s_path
                    break
        
        # If it's a reference page, we prefer extracting an embedded screenshot if possible
        if img_path and 'reference_pages' in img_path:
            resolved_path = None
            if embedded_screenshots:
                # 1. Try to find the best match based on context text
                best_ss = find_best_embedded_screenshot(context_text, embedded_screenshots) if context_text else None
                if best_ss and os.path.exists(best_ss):
                    resolved_path = best_ss
                else:
                    # 2. Fallback: Check if there are embedded screenshots belonging to this PDF page number
                    if slide_num:
                        page_screenshots = [ss for ss in embedded_screenshots if ss.get('page_number') == slide_num]
                        if page_screenshots:
                            # If only one image on that page, use it
                            if len(page_screenshots) == 1:
                                resolved_path = page_screenshots[0].get('image_path')
                            else:
                                # If multiple, run context matching restricted to page screenshots if context exists
                                best_page_ss = find_best_embedded_screenshot(context_text, page_screenshots) if context_text else None
                                if best_page_ss:
                                    resolved_path = best_page_ss
                                else:
                                    # Default to the first embedded screenshot from that page
                                    resolved_path = page_screenshots[0].get('image_path')
            img_path = resolved_path
    else:
        if ts and ts in timestamp_to_frame:
            frame_fname = timestamp_to_frame[ts]
            img_path = os.path.join('screenshots', frame_fname)
        else:
            fallback_name = f"{block_id}_{ts.replace(':', '')}.jpg"
            img_path = f"screenshots/{fallback_name}"
            if fallback_name in frames:
                frame_fname = fallback_name
            elif f"{fallback_name}*" in frames:
                frame_fname = f"{fallback_name}*"
    return img_path, frame_fname

def insert_image_for_vm(doc, vm, block_id, slides, timestamp_to_frame, frames, inserted_filenames, inserted_ocrs, embedded_screenshots=None, context_text=""):
    img_path, frame_fname = get_vm_image_path(vm, block_id, slides, timestamp_to_frame, frames, embedded_screenshots, context_text)
    if img_path and os.path.exists(img_path):
        current_ocr = ""
        if frame_fname and frame_fname in frames:
            current_ocr = frames[frame_fname].get('ocr_text', '')
        
        is_duplicate = os.path.basename(img_path) in inserted_filenames
        is_slide = "slides" in img_path or "slide_" in os.path.basename(img_path)
        
        current_hash = None
        try:
            current_hash = imagehash.dhash(Image.open(img_path))
        except Exception as e:
            logging.warning(f"Could not compute imagehash for {img_path}: {e}")

        if current_hash is not None:
            for item in inserted_ocrs:
                prev_hash = item.get('hash')
                prev_ocr = item.get('ocr', '')
                
                # Check visual difference (lower threshold to 3 to ensure distinct steps are kept)
                if prev_hash is not None:
                    diff = current_hash - prev_hash
                    if diff <= 3:
                        # Visually similar: check if OCR is also similar
                        if current_ocr.strip() and prev_ocr.strip():
                            if are_ocr_texts_similar(current_ocr, prev_ocr, threshold=0.95):
                                is_duplicate = True
                                logging.info(f"Duplicate detected: visually similar (diff {diff} <= 3) and OCR similar (> 0.95)")
                                break
                        else:
                            is_duplicate = True
                            logging.info(f"Duplicate detected: visually similar (diff {diff} <= 3) with empty OCR")
                            break
                else:
                    # Fallback to OCR check only
                    if current_ocr.strip() and prev_ocr.strip():
                        if are_ocr_texts_similar(current_ocr, prev_ocr, threshold=0.95):
                            is_duplicate = True
                            logging.info(f"Duplicate detected: OCR similar (> 0.95) and no hash available")
                            break
        
        if is_duplicate or is_logo_frame(current_ocr):
            logging.info(f"Skipping duplicate or logo slide image: {frame_fname or img_path}")
            return False
        else:
            if add_image_if_exists(doc, img_path):
                inserted_filenames.add(os.path.basename(img_path))
                inserted_ocrs.append({
                    'hash': current_hash,
                    'ocr': current_ocr
                })
                return True
    return False

def build_document(concept_map_path, frame_manifest_path, slide_manifest_path, output_path, audit_feedback_path=None):
    global CLOZE_ANSWERS
    CLOZE_ANSWERS = []
    
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    # Do NOT set explicit dark color on Normal style (leave it as Auto) so body text is dark-mode auto-invertible

    # Load manifests
    if not os.path.exists(concept_map_path):
        logging.error(f"Error: Concept block map not found at {concept_map_path}")
        return False, None

    with open(concept_map_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
        lecture_title_from_map = ""
        if isinstance(data, dict):
            lecture_title_from_map = data.get("lecture_title", "")
            concept_blocks = data.get("blocks", [])
        else:
            concept_blocks = data

    # Guard against empty concept_blocks list
    if not concept_blocks or not isinstance(concept_blocks, list) or len(concept_blocks) == 0:
        logging.warning("Warning: Concept block map is empty. Generating minimal document.")
        concept_blocks = []

    frames = {}
    if frame_manifest_path and os.path.exists(frame_manifest_path):
        with open(frame_manifest_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if isinstance(data, dict) and "frames" in data and isinstance(data["frames"], list):
                for item in data["frames"]:
                    fname = item.get("filename")
                    if fname:
                        frames[fname] = item
            elif isinstance(data, list):
                for item in data:
                    fname = item.get("filename")
                    if fname:
                        frames[fname] = item
            else:
                frames = data

    slides = []
    if slide_manifest_path and os.path.exists(slide_manifest_path):
        with open(slide_manifest_path, 'r', encoding='utf-8') as f:
            slides = json.load(f)
            
    reference_manifest_path = "reference_manifest.json"
    if os.path.exists(reference_manifest_path):
        try:
            with open(reference_manifest_path, 'r', encoding='utf-8') as f:
                ref_slides = json.load(f)
                if isinstance(ref_slides, list):
                    slides.extend(ref_slides)
        except Exception as e:
            logging.warning(f"Could not load reference_manifest.json: {e}")
            
    embedded_screenshots = []
    embedded_manifest_path = "embedded_manifest.json"
    if os.path.exists(embedded_manifest_path):
        try:
            with open(embedded_manifest_path, 'r', encoding='utf-8') as f:
                embedded_screenshots = json.load(f)
        except Exception as e:
            logging.warning(f"Could not load embedded_manifest.json: {e}")

    # SEMANTIC COHERENCE CHECK: Warn if concept blocks and frames appear mismatched
    import re  # ensure re is always in scope for this function
    if concept_blocks and frames:
        concept_keywords = set()
        for block in concept_blocks[:3]:  # Check first few blocks
            # Titles
            title = block.get('title', '').lower()
            concept_keywords.update(re.findall(r'\b[a-z0-9_\-\+]{1,}\b', title))
            
            # Explanations
            expl = block.get('explanation', '').lower()
            concept_keywords.update(re.findall(r'\b[a-z0-9_\-\+]{1,}\b', expl))
            
            # Concept terms
            for c in block.get('concepts', []):
                if isinstance(c, dict):
                    term = c.get('term', '').lower()
                else:
                    term = str(c).split(':', 1)[0].lower()
                concept_keywords.update(re.findall(r'\b[a-z0-9_\-\+]{1,}\b', term))
                
            # Example sentences
            for ex in block.get('examples', []):
                sent = ex.get('sentence', '').lower()
                concept_keywords.update(re.findall(r'\b[a-z0-9_\-\+]{1,}\b', sent))
        
        # Remove common noise words and short grammatical/filler words
        noise = {
            "about", "above", "after", "again", "against", "along", "could", "would", "should", 
            "their", "there", "these", "those", "which", "where", "under", "family",
            "the", "a", "an", "and", "or", "but", "if", "then", "else", "of", "at", "by", "for", 
            "with", "to", "from", "in", "on", "into", "over", "is", "was", "were", "are", "be", 
            "been", "have", "has", "had", "do", "does", "did", "we", "you", "he", "she", "it", 
            "they", "i", "my", "our", "your", "his", "her", "its", "them", "us", "this", "that", 
            "not", "no", "yes", "so", "as", "than", "very", "can", "will", "just", "now"
        }
        concept_keywords = concept_keywords - noise
        
        frame_keywords = set()
        for fname, info in list(frames.items())[:10]:  # Check first 10 frames
            ocr = info.get('ocr_text', '').lower()
            words = re.findall(r'\b[a-z0-9_\-\+]{1,}\b', ocr)
            frame_keywords.update(words)
        frame_keywords = frame_keywords - noise
        
        common = concept_keywords & frame_keywords
        overlap_ratio = len(common) / max(len(concept_keywords), 1)
        
        if overlap_ratio < 0.15:
            logging.warning(f"⚠️  SEMANTIC COHERENCE WARNING: Concept blocks and frames may be from different lectures!")
            logging.warning(f"   Concept keywords: {list(concept_keywords)[:15]}")
            logging.warning(f"   Frame keywords: {list(frame_keywords)[:15]}")
            logging.warning(f"   Overlap ratio: {overlap_ratio:.2f} (expected > 0.15)")
        else:
            logging.info(f"✅ Semantic coherence check passed. Overlap ratio: {overlap_ratio:.2f}")

    # Build a lookup from timestamp to frame filename
    timestamp_to_frame = {}
    for filename, info in frames.items():
        ts = info.get('timestamp')
        if ts:
            ts_clean = ts.rstrip('*')
            timestamp_to_frame[ts_clean] = filename

    # Title
    lecture_title = lecture_title_from_map or "Lecture Reconstruction Notes"
    if concept_blocks and isinstance(concept_blocks, list) and len(concept_blocks) > 0:
        if not lecture_title_from_map and "lecture_title" in concept_blocks[0]:
            lecture_title = concept_blocks[0]["lecture_title"]
        elif not lecture_title_from_map:
            first_title = concept_blocks[0].get('title', '')
            if first_title and not re.match(r'^.*\(Questions?\s*\d+', first_title):
                lecture_title = first_title
    add_custom_heading(doc, lecture_title, 0)

    # Section 1 & 2 Outline headers removed to follow single-subject outline policy

    inserted_ocrs = []
    inserted_filenames = set()
    global_callout_count = 0
    MAX_CALLOUTS = 6
    
    total_callouts_rendered = 0
    enforce_callout_cap = False
    if audit_feedback_path and os.path.exists(audit_feedback_path):
        try:
            with open(audit_feedback_path, 'r', encoding='utf-8') as f:
                audit_feedback = json.load(f)
                if "24" in audit_feedback:  # Gate 24: Total Callout Box Cap
                    enforce_callout_cap = True
        except Exception as e:
            logging.warning(f"Could not load audit feedback: {e}")
    global_callout_count = 0
    MAX_CALLOUTS = 6
    for block in concept_blocks:
        block_id = block.get('block_id', 'CB')
        title = block.get('title', 'Untitled Concept')
        # Clean "Concept Block X: " prefix from title using regex
        clean_title = re.sub(r'(?i)^\s*Concept\s*Block\s*\d*\s*[:\-]?\s*', '', title).strip()
        add_custom_heading(doc, clean_title, level=2)


        # Check if the block has a custom flow
        flow = block.get('flow', [])
        visual_moments = block.get('visual_moments', [])
        inserted_vm_indices = set()
        seen_rules_in_block = []

        if flow:
            # Loop through elements in custom flow
            for elem in flow:
                el_type = elem.get('type')
                if el_type == 'paragraph':
                    add_formatted_explanation_paragraphs(doc, elem.get('text', ''))
                elif el_type == 'concept':
                    p_c = doc.add_paragraph(style='List Bullet')
                    term = elem.get('term', '')
                    definition = elem.get('definition', '')
                    add_bold_prefix(p_c, term + ": ")
                    add_rich_runs(p_c, definition)
                elif el_type == 'example':
                    ex_idx = elem.get('index')
                    ex = None
                    examples = block.get('examples', [])
                    if ex_idx is not None and ex_idx < len(examples):
                        ex = examples[ex_idx]
                    else:
                        target_sent = elem.get('sentence', '')
                        for e in examples:
                            if e.get('sentence', '').strip() == target_sent.strip():
                                ex = e
                                break
                    if ex:
                        wrong = ex.get('wrong')
                        correct = ex.get('correct')
                        is_question = False
                        if wrong and correct:
                            add_correction(doc, wrong, correct)
                        else:
                            p_q = doc.add_paragraph()
                            sentence = ex.get('scenario_or_problem', ex.get('sentence', 'Example'))
                            is_question = sentence.endswith('?') or sentence.startswith(('Q:', 'Explain', 'Describe', 'Why', 'How', 'What', 'Define', 'List', 'Solve', 'Calculate', 'Find', 'Determine', 'Evaluate', 'Compare', 'Verify'))
                            pref = "Q: " if is_question else "Example: "
                            add_bold_prefix(p_q, pref)
                            add_rich_runs(p_q, sentence)

                        rule = ex.get('core_principles', ex.get('rule', ''))
                        is_redundant_rule = False
                        if rule:
                            is_redundant_rule = any(are_rules_similar(rule, prev) for prev in seen_rules_in_block)
                            if working_text and (rule.lower() in working_text.lower() or working_text.lower() in rule.lower()):
                                is_redundant_rule = True
                            if sentence and (rule.lower() in sentence.lower() or sentence.lower() in rule.lower()):
                                is_redundant_rule = True
                                
                            if not is_redundant_rule:
                                p_rule = doc.add_paragraph()
                                label = "Rule: " if is_question else ""
                                if label:
                                    add_bold_prefix(p_rule, label)
                                add_rich_runs(p_rule, f"<b>{rule}</b>" if not label else rule)
                                seen_rules_in_block.append(rule)

                        working_text = ex.get('step_by_step_logic', ex.get('working', ''))
                        if working_text:
                            # Drop the bulky header and just render the steps directly
                            add_formatted_explanation_paragraphs(doc, working_text)

                        # Render method2 if present
                        method2 = ex.get('method2', elem.get('method2'))
                        if method2:
                            p_m2 = doc.add_paragraph()
                            p_m2.paragraph_format.space_before = Pt(4)
                            p_m2.paragraph_format.space_after = Pt(4)
                            add_bold_prefix(p_m2, "Alternative Approach (Method 2): ")
                            add_rich_runs(p_m2, method2)



                        ans_text = get_explicit_answer(ex, working_text)
                        if ans_text:
                            p_ans = doc.add_paragraph()
                            add_bold_prefix(p_ans, "Answer: ")
                            add_rich_runs(p_ans, ans_text)

                        # Render corresponding visual moment inline
                        target_ts = ex.get('timestamp')
                        vm_to_insert = None
                        target_idx = -1
                        if target_ts:
                            for v_idx, vm in enumerate(visual_moments):
                                if vm.get('timestamp') == target_ts:
                                    vm_to_insert = vm
                                    target_idx = v_idx
                                    break
                        if not vm_to_insert:
                            try:
                                ex_idx_in_list = examples.index(ex)
                            except ValueError:
                                ex_idx_in_list = -1
                            if ex_idx_in_list != -1 and ex_idx_in_list < len(visual_moments) and ex_idx_in_list not in inserted_vm_indices:
                                vm_to_insert = visual_moments[ex_idx_in_list]
                                target_idx = ex_idx_in_list
                            else:
                                for v_idx in range(len(visual_moments)):
                                    if v_idx not in inserted_vm_indices:
                                        vm_to_insert = visual_moments[v_idx]
                                        target_idx = v_idx
                                        break
                        if vm_to_insert:
                            context_text = ex.get('sentence', '') + " " + ex.get('working', '')
                            success_inserted = insert_image_for_vm(doc, vm_to_insert, block_id, slides, timestamp_to_frame, frames, inserted_filenames, inserted_ocrs, embedded_screenshots, context_text)
                            if target_idx != -1:
                                inserted_vm_indices.add(target_idx)
                elif el_type == 'trap':
                    if global_callout_count >= MAX_CALLOUTS: continue
                    add_trap(doc, elem.get('text', ''))
                    global_callout_count += 1
                    total_callouts_rendered += 1
                elif el_type == 'trick':
                    if global_callout_count >= MAX_CALLOUTS: continue
                    add_trick(doc, elem.get('text', ''))
                    global_callout_count += 1
                    total_callouts_rendered += 1
                elif el_type == 'quote':
                    if global_callout_count >= MAX_CALLOUTS: continue
                    add_quote(doc, elem.get('text', ''))
                    global_callout_count += 1
                    total_callouts_rendered += 1
                elif el_type == 'highlight':
                    p_i = doc.add_paragraph()
                    add_bold_prefix(p_i, "⭐ [IMPORTANT] ")
                    add_rich_runs(p_i, elem.get('text', ''))
                elif el_type == 'table':
                    headers = elem.get('headers', [])
                    rows = elem.get('rows', [])
                    if headers and rows:
                        table_title = elem.get('title', 'Reference Table')
                        add_custom_heading(doc, table_title, level=3)
                        add_styled_table(doc, headers, rows)
                        doc.add_paragraph()
                elif el_type == 'cornell_block':
                    cue = elem.get('cue', '')
                    content = elem.get('content', [])
                    srs_tag = elem.get('srs_tag')
                    add_cornell_block(doc, cue, content, srs_tag)
        else:
            # Fallback to standard sequential rendering

            

            
            if title == "Super Key":
                add_super_key_additions(doc)

            # Mnemonic section
            mnemonic = block.get('mnemonic')
            if mnemonic:
                p_m_lbl = doc.add_paragraph()
                p_m_lbl.paragraph_format.space_before = Pt(6)
                p_m_lbl.paragraph_format.space_after = Pt(2)
                add_bold_prefix(p_m_lbl, "@ Mnemonic:")
                
                p_m_val = doc.add_paragraph()
                p_m_val.paragraph_format.space_before = Pt(0)
                p_m_val.paragraph_format.space_after = Pt(4)
                add_rich_runs(p_m_val, mnemonic.get('text', ''))
                
                p_bd_lbl = doc.add_paragraph()
                p_bd_lbl.paragraph_format.space_before = Pt(4)
                p_bd_lbl.paragraph_format.space_after = Pt(2)
                add_bold_prefix(p_bd_lbl, "The Breakdown:")
                
                for item in mnemonic.get('breakdown', []):
                    p_item = doc.add_paragraph()
                    p_item.paragraph_format.space_before = Pt(0)
                    p_item.paragraph_format.space_after = Pt(2)
                    add_rich_runs(p_item, item)
                    
                p_sp_lbl = doc.add_paragraph()
                p_sp_lbl.paragraph_format.space_before = Pt(4)
                p_sp_lbl.paragraph_format.space_after = Pt(2)
                add_bold_prefix(p_sp_lbl, "The \"Specialized\" Trio (The remaining 3):")
                
                for item in mnemonic.get('specialized', []):
                    p_item = doc.add_paragraph()
                    p_item.paragraph_format.space_before = Pt(0)
                    p_item.paragraph_format.space_after = Pt(2)
                    add_rich_runs(p_item, item)
                doc.add_paragraph()

            # Table data rendering
            table_data = block.get('table')
            if table_data:
                headers = table_data.get('headers', [])
                rows = table_data.get('rows', [])
                if headers and rows:
                    table_title = table_data.get('title', 'Reference Table')
                    add_custom_heading(doc, table_title, level=3)
                    add_styled_table(doc, headers, rows)
                    doc.add_paragraph()

            # Key Concepts and Definitions
            concepts = block.get('concepts', [])
            if concepts:
                add_custom_heading(doc, "Key Concepts & Definitions:", level=3)
                for item in concepts:
                    p_c = doc.add_paragraph(style='List Bullet')
                    if isinstance(item, dict):
                        term = item.get('term', '')
                        definition = item.get('definition', '')
                    else:
                        parts = str(item).split(':', 1)
                        term = parts[0].strip()
                        definition = parts[1].strip() if len(parts) > 1 else ''
                    add_bold_prefix(p_c, term + ": ")
                    add_rich_runs(p_c, definition)

            # Exam Imp section
            exam_imp = block.get('exam_imp', [])
            if exam_imp:
                p_head = doc.add_paragraph()
                p_head.paragraph_format.space_before = Pt(6)
                p_head.paragraph_format.space_after = Pt(2)
                run_head = p_head.add_run("EXAM IMP:")
                run_head.bold = True
                run_head.underline = True
                run_head.font.size = Pt(10)
                run_head.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
                run_head.font.name = 'Calibri'
                
                for item in exam_imp:
                    p_item = doc.add_paragraph()
                    p_item.paragraph_format.space_before = Pt(0)
                    p_item.paragraph_format.space_after = Pt(3)
                    add_rich_runs(p_item, item)
                doc.add_paragraph()

            # Examples & Illustrations
            examples = block.get('examples', [])
            solved_examples = []
            hw_examples = []
            if examples:
                solved_examples = [ex for ex in examples if not ex.get('is_homework')]
                hw_examples = [ex for ex in examples if ex.get('is_homework')]

            if solved_examples:
                add_custom_heading(doc, "Examples & Illustrations:", level=3)
                for idx, ex in enumerate(solved_examples):
                    wrong = ex.get('wrong')
                    correct = ex.get('correct')
                    is_question = False
                    if wrong and correct:
                        add_correction(doc, wrong, correct)
                    else:
                        sentence = ex.get('scenario_or_problem', ex.get('sentence', 'Example'))
                        is_question = sentence.endswith('?') or sentence.startswith(('Q:', 'Explain', 'Describe', 'Why', 'How', 'What', 'Define', 'List', 'Solve', 'Calculate', 'Find', 'Determine', 'Evaluate', 'Compare', 'Verify'))
                        pref = "Q: " if is_question else "Example: "

                    rule = ex.get('core_principles', ex.get('rule', ''))
                    is_redundant_rule = False
                    if rule:
                        is_redundant_rule = any(are_rules_similar(rule, prev) for prev in seen_rules_in_block)

                    working_text = ex.get('step_by_step_logic', ex.get('working', ''))
                    cloze_text = ex.get('cloze_text')
                    method2 = ex.get('method2')

                    ans_text = get_explicit_answer(ex, working_text)

                    # Render Cornell Block if cues are present (Layer 2 of Friction-Optimized Note Matrix)
                    cornell_cues = ex.get('cornell_cues', [])
                    srs_tag = ex.get('srs_tag')
                    if cornell_cues:
                        cue_text = "\n\n".join(cornell_cues)
                        content_list = []
                        content_list.append(f"<b>{pref}</b> {sentence}")
                        if rule and not is_redundant_rule:
                            label = "Applicable Rule: " if is_question else "Key Concept: "
                            content_list.append(f"<b>{label}</b> {rule}")
                            seen_rules_in_block.append(rule)
                        if cloze_text:
                            content_list.append(f"<b>Explanation (Active Recall):</b> {cloze_text}")
                        elif working_text:
                            label_work = "Explanation/Working:" if is_question else "Explanation:"
                            content_list.append(f"<b>{label_work}</b> {working_text}")
                        if method2:
                            content_list.append(f"<b>Alternative Approach (Method 2):</b> {method2}")

                        if ans_text:
                            content_list.append(f"<b>Answer:</b> {ans_text}")
                        add_cornell_block(doc, cue_text, content_list, srs_tag)
                    else:
                        # Fallback to standard sequential rendering
                        if wrong and correct:
                            pass # Already handled correction
                        else:
                            p_q = doc.add_paragraph()
                            add_bold_prefix(p_q, pref)
                            add_rich_runs(p_q, sentence)

                        if rule and not is_redundant_rule:
                            # aggressively check if rule is just repeating the sentence or working text
                            if working_text and (rule.lower() in working_text.lower() or working_text.lower() in rule.lower()):
                                is_redundant_rule = True
                            if sentence and (rule.lower() in sentence.lower() or sentence.lower() in rule.lower()):
                                is_redundant_rule = True
                                
                            if not is_redundant_rule:
                                p_rule = doc.add_paragraph()
                                label = "Rule: " if is_question else ""
                                if label:
                                    add_bold_prefix(p_rule, label)
                                add_rich_runs(p_rule, f"<b>{rule}</b>" if not label else rule)
                                seen_rules_in_block.append(rule)

                        if cloze_text:
                            p_cl = doc.add_paragraph()
                            add_bold_prefix(p_cl, "Explanation (Active Recall): ")
                            add_rich_runs(p_cl, cloze_text)
                        elif working_text:
                            # Drop the bulky 'Explanation:' header and just render the steps
                            add_formatted_explanation_paragraphs(doc, working_text)

                        # Render method2 if present
                        if method2:
                            p_m2 = doc.add_paragraph()
                            p_m2.paragraph_format.space_before = Pt(4)
                            p_m2.paragraph_format.space_after = Pt(4)
                            add_bold_prefix(p_m2, "Alternative Approach (Method 2): ")
                            add_rich_runs(p_m2, method2)



                        if ans_text:
                            p_ans = doc.add_paragraph()
                            add_bold_prefix(p_ans, "Answer: ")
                            add_rich_runs(p_ans, ans_text)

                    # Place corresponding visual moment inline right under the example
                    vm_to_insert = None
                    target_ts = ex.get('timestamp')
                    target_idx = -1
                    if target_ts:
                        for v_idx, vm in enumerate(visual_moments):
                            if vm.get('timestamp') == target_ts:
                                vm_to_insert = vm
                                target_idx = v_idx
                                break
                    if vm_to_insert:
                        context_text = sentence + " " + (cloze_text or working_text)
                        insert_image_for_vm(doc, vm_to_insert, block_id, slides, timestamp_to_frame, frames, inserted_filenames, inserted_ocrs, embedded_screenshots, context_text)
                        if target_idx != -1:
                            inserted_vm_indices.add(target_idx)

            if hw_examples:
                add_custom_heading(doc, "Homework Questions (HW Que): Try:", level=3)
                for idx, ex in enumerate(hw_examples):
                    p_q = doc.add_paragraph()
                    sentence = ex.get('sentence', 'Example')
                    add_bold_prefix(p_q, "Q: ")
                    add_rich_runs(p_q, sentence)

                    rule = ex.get('core_principles', ex.get('rule', ''))
                    is_redundant_rule = False
                    if rule:
                        if working_text and (rule.lower() in working_text.lower() or working_text.lower() in rule.lower()):
                            is_redundant_rule = True
                        if sentence and (rule.lower() in sentence.lower() or sentence.lower() in rule.lower()):
                            is_redundant_rule = True
                        
                        if not is_redundant_rule:
                            p_rule = doc.add_paragraph()
                            add_bold_prefix(p_rule, "Rule: ")
                            add_rich_runs(p_rule, rule)

                    working_text = ex.get('step_by_step_logic', ex.get('working', ''))
                    if working_text:
                        add_formatted_explanation_paragraphs(doc, working_text)

                    # Render method2 if present
                    method2 = ex.get('method2')
                    if method2:
                        p_m2 = doc.add_paragraph()
                        p_m2.paragraph_format.space_before = Pt(4)
                        p_m2.paragraph_format.space_after = Pt(4)
                        add_bold_prefix(p_m2, "Alternative Approach (Method 2): ")
                        add_rich_runs(p_m2, method2)



                    ans_text = get_explicit_answer(ex, working_text)
                    if ans_text:
                        p_ans = doc.add_paragraph()
                        add_bold_prefix(p_ans, "Answer: ")
                        add_rich_runs(p_ans, ans_text)

            # Important Points / Teacher's Emphasis
            important_points = block.get('important_points', [])

            if important_points:
                add_custom_heading(doc, "Key Highlights (⭐ Teacher's Emphasis):", level=3)
                for item in important_points:
                    p_i = doc.add_paragraph()
                    add_bold_prefix(p_i, "⭐ [IMPORTANT] ")
                    add_rich_runs(p_i, item)

            # Exercise Questions
            exercises = block.get('exercise_questions', [])
            if exercises:
                real_exercises = [eq for eq in exercises if isinstance(eq, str) and eq.strip()]
                if real_exercises:
                    p = doc.add_paragraph()
                    add_bold_prefix(p, "Homework Questions (HW Que): Try:")
                    for eq in real_exercises:
                        p_ex = doc.add_paragraph(style='List Bullet')
                        add_rich_runs(p_ex, eq)

            # Boundary Questions — SKIPPED (removed to reduce verbosity)
            # boundary_qs = block.get('boundary_questions', [])

            # Visual Moments & Images (Leftovers)
            leftover_moments = [vm for i, vm in enumerate(visual_moments) if i not in inserted_vm_indices]
            for vm in leftover_moments:
                context_text = block.get('explanation', '')
                insert_image_for_vm(doc, vm, block_id, slides, timestamp_to_frame, frames, inserted_filenames, inserted_ocrs, embedded_screenshots, context_text)

            # Render block-level elements with document-wide caps
            block_quotes = block.get('teacher_quotes', [])
            for q in block_quotes[:1]:  # Cap at 1 teacher quote per block
                if global_callout_count >= MAX_CALLOUTS:
                    break
                add_quote(doc, q)
                global_callout_count += 1
                total_callouts_rendered += 1
            # Support both old schema (traps/tricks) and new schema (teacher_cautions)
            traps = block.get('traps') or []
            tricks = block.get('tricks') or []
            cautions = (block.get('teacher_cautions') or []) or (traps[:1] + tricks[:1])
            for c in cautions:
                if global_callout_count >= MAX_CALLOUTS:
                    break
                if c and isinstance(c, str) and c.strip():
                    add_trap(doc, c)
                    global_callout_count += 1
                    total_callouts_rendered += 1





        doc.add_paragraph()

    # Load profile
    profile = {}
    if os.path.exists("lecture_profile.json"):
        try:
            with open("lecture_profile.json", "r", encoding="utf-8") as f:
                profile = json.load(f)
        except Exception as e:
            logging.warning(f"Could not load lecture_profile.json: {e}")

    # Section 3 and Section 4 REMOVED — they redundantly re-printed concepts and rules
    # already present in the main body. Revision boxes at the end of each section serve
    # the same purpose without adding word count.

    # Relocated: Add a single Quick Revision Summary box at the very end of the document
    rev_bullets = []
    for idx, block in enumerate(concept_blocks):
        b_num = idx + 1
        b_title = block.get('title', f"Block {b_num}")
        b_title_clean = re.sub(r'(?i)^\s*Concept\s*Block\s*\d*\s*[:\-]?\s*', '', b_title).strip()
        
        traps = block.get('traps') or []
        cautions = (block.get('teacher_cautions') or []) or traps
        
        rules = []
        for ex in block.get('examples', []):
            rule_text = ex.get('rule') or ex.get('core_principles')
            if rule_text:
                rules.append(rule_text)
                
        if cautions and len(cautions) > 0:
            rev_bullets.append(f"⭐ **{b_title_clean} (Warning)**: {cautions[0]}")
        elif rules and len(rules) > 0:
            rev_bullets.append(f"💡 **{b_title_clean} (Rule)**: {rules[0]}")
        else:
            rev_bullets.append(f"📌 **{b_title_clean}**: Refer to detailed notes.")
            
    if rev_bullets:
        add_custom_heading(doc, "Quick Revision Summary", level=1)
        add_revision_box(doc, rev_bullets)
        doc.add_paragraph()

    # Section 5: Appendix: Cloze Deletion Answer Key
    if CLOZE_ANSWERS:
        add_custom_heading(doc, "Appendix: Cloze Deletion Answer Key", level=1)
        headers = ["Cloze ID", "Correct Answer", "Teacher's Hint / Context"]
        rows = []
        for item in CLOZE_ANSWERS:
            cid = f"[{item['id']}]"
            ans = f"<b>{item['answer']}</b>"
            hnt = item['hint'] or "No hint provided"
            rows.append([cid, ans, hnt])
        add_styled_table(doc, headers, rows)

    import datetime
    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    
    doc.save(output_path)
    logging.info(f"Notes document generated successfully at: {output_path}")

    # Save inserted image metadata only after the target DOCX exists so audit can
    # verify the image list belongs to the same file instead of trusting stale data.
    try:
        payload = {
            "_docx": build_docx_metadata(output_path),
            "images": list(inserted_filenames),
        }
        with open("inserted_images.json", "w", encoding="utf-8") as f:
            json.dump(payload, f, indent=2)
        logging.info(f"Saved {len(inserted_filenames)} inserted image names to inserted_images.json")
    except Exception as e:
        logging.warning(f"Could not save inserted_images.json: {e}")

    return True, None

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Generate Word document notes from sub-agent manifests.")
    parser.add_argument('--concept-map', default='concept_block_map.json')
    parser.add_argument('--frame-manifest', default='frame_manifest.json')
    parser.add_argument('--slide-manifest', default='slide_manifest.json')
    parser.add_argument('--output', default='notes-output/LECTURE_NOTES.docx')
    parser.add_argument('--audit-feedback', default=None, help='Path to audit feedback JSON')
    args = parser.parse_args()
    success, archive_path = build_document(args.concept_map, args.frame_manifest, args.slide_manifest, args.output, args.audit_feedback)
    import sys
    sys.exit(0 if success else 1)
