import glob
#!/usr/bin/env python3
import os
import json
import argparse
import sys
import re
import logging
from docx import Document
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from scripts.generate_docx import format_math_text, clean_attributions

audit_feedback = {}

class FeedbackHandler(logging.Handler):
    def emit(self, record):
        msg = self.format(record)
        match = re.search(r'\[FAIL\] Gate\s+(\d+):\s*(.*)', msg)
        if match:
            gate_num = match.group(1)
            reason = match.group(2)
            if gate_num not in audit_feedback:
                audit_feedback[gate_num] = []
            audit_feedback[gate_num].append(reason)

# Configure logging
os.makedirs("logs", exist_ok=True)
feedback_handler = FeedbackHandler()
feedback_handler.setLevel(logging.WARNING)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[
        logging.FileHandler("logs/pipeline.log"),
        logging.StreamHandler(sys.stderr),
        feedback_handler
    ]
)

def count_worked_examples_in_docx(doc):
    prefixes = ('Q:', 'Example:', 'Situation:', 'Examples:', 'Illustration:', 'Wrong:')
    return sum(1 for p in get_all_paragraphs(doc) if any(p.text.strip().startswith(pref) for pref in prefixes))

def get_all_text_from_docx(doc):
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    texts.append(p.text)
    return " ".join(texts)

def get_all_paragraphs(doc):
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs

def is_paragraph_in_table(paragraph):
    elem = paragraph._p
    while elem is not None:
        if elem.tag.endswith('}tc'):
            return True
        elem = elem.getparent()
    return False

def calculate_friction_index(doc, concept_blocks):

    all_doc_text = get_all_text_from_docx(doc)
    
    clean_text = re.sub(r'[^\w\s]', '', all_doc_text)
    total_words = len(clean_text.split())
    if total_words == 0:
        return 0.0, 0, 0, 0
        
    import json
    cloze_count = json.dumps(concept_blocks).count('<cloze')
    
    cornell_cues_count = 0
    for b in concept_blocks:
        flow = b.get('flow', [])
        for elem in flow:
            if elem.get('type') == 'cornell_block':
                cornell_cues_count += 1
                
    friction_index = (cloze_count + cornell_cues_count) / total_words
    return friction_index, total_words, cloze_count, cornell_cues_count

def build_docx_metadata(docx_path):
    abs_path = os.path.abspath(docx_path)
    metadata = {"path": abs_path}
    try:
        stat = os.stat(docx_path)
        metadata["size"] = stat.st_size
        metadata["mtime"] = int(stat.st_mtime)
    except OSError:
        metadata["size"] = None
        metadata["mtime"] = None
    return metadata

def inserted_images_match_docx(inserted_payload, docx_metadata):
    if not isinstance(inserted_payload, dict):
        return False
    payload_docx = inserted_payload.get("_docx")
    if not isinstance(payload_docx, dict):
        return False
    return (
        payload_docx.get("path") == docx_metadata.get("path")
        and payload_docx.get("size") == docx_metadata.get("size")
        and payload_docx.get("mtime") == docx_metadata.get("mtime")
    )

def run_audit(docx_path, concept_map_path, frame_manifest_path, slide_manifest_path):
    logging.info(f"Starting 24-Gate Quality Audit on: {docx_path}")
    
    failed_gates = {
        'Gate 1: Structural Integrity':    False,
        'Gate 2: Revision Box Placement':  False,
        'Gate 3: Chronological Flow':      False,
        'Gate 4: Content Completeness':    False,
        'Gate 5: Factual Accuracy':        False,
        'Gate 6: Image Integrity':         False,
        'Gate 7: Minimum Counts':          False,
        'Gate 8: Source Traceability':     False,
        'Gate 9: Slide Handling':          False,
        'Gate 10: Example Coverage':       False,
        'Gate 11: Visual Coverage':        False,
        'Gate 12: Exercise Content':       False,
        'Gate 13: Quote Quality':          False,
        'Gate 14: Meaningful Titles':      False,
        'Gate 15: Explanation Conciseness': False,
        'Gate 16: Table Presence':          False,
        'Gate 17: Sequence Integrity':      False,
        'Gate 18: Exact Worked Examples':  False,
        'Gate 19: Friction Index Constraint': False,
        'Gate 20: Transcript Coverage':     False,
        'Gate 21: English Enforcement':     False,
        'Gate 22: Styling and Highlighting Conformity': False,
        'Gate 23: Word Count Budget':       False,
        'Gate 24: Callout Box Cap':         False,
        'Gate 25: Short Note Audit':        False,
    }

    if not os.path.exists(docx_path):
        logging.error(f"[FAIL] Target document not found at: {docx_path}")
        return False, failed_gates
        
    try:
        doc = Document(docx_path)
    except Exception as e:
        logging.error(f"[FAIL] Failed to open document {docx_path}: {e}")
        return False, failed_gates

    concept_blocks = []
    concept_map_title = ""
    if os.path.exists(concept_map_path):
        try:
            with open(concept_map_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if isinstance(data, dict):
                    concept_map_title = data.get("lecture_title", "")
                    concept_blocks = data.get("blocks", [])
                else:
                    concept_blocks = data
                    if concept_blocks and isinstance(concept_blocks[0], dict):
                        concept_map_title = concept_blocks[0].get("lecture_title", "")
        except Exception as e:
            logging.error(f"Failed to read concept map: {e}")

    frames = {}
    if os.path.exists(frame_manifest_path):
        try:
            with open(frame_manifest_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if isinstance(data, dict) and "frames" in data and isinstance(data["frames"], list):
                    for item in data["frames"]:
                        fname = item.get("filename")
                        if fname:
                            frames[fname] = item
                elif isinstance(data, list):
                    for item in data:
                        fname = item.get("filename")
                        if fname:
                            frames[fname] = item
                else:
                    frames = data
        except Exception as e:
            logging.error(f"Failed to read frame manifest: {e}")

    slides = []
    if os.path.exists(slide_manifest_path):
        try:
            with open(slide_manifest_path, 'r', encoding='utf-8') as f:
                slides = json.load(f)
        except Exception as e:
            logging.error(f"Failed to read slide manifest: {e}")

    h2 = rev = trap = trick = quote = vis_fail = attr_fail = 0
    banned = [
        "the lecturer says", "the teacher explains", "the instructor mentions",
        "this is discussed in the lecture", "the teacher describes",
        "the teacher outlines", "the teacher demonstrates", "the teacher analyzes",
        "the teacher shares", "the teacher introduces", "the teacher reviews",
        "the teacher teaches", "the teacher shows", "the teacher discusses",
        "the instructor outlines", "the instructor demonstrates", "the instructor analyzes",
        "the instructor shares", "the instructor introduces", "the instructor reviews",
        "the instructor teaches", "the instructor shows", "the instructor discusses",
        "the lecturer outlines", "the lecturer demonstrates", "the lecturer analyzes",
        "the lecturer shares", "the lecturer introduces", "the lecturer reviews",
        "the lecturer teaches", "the lecturer shows", "the lecturer discusses",
        "we see", "we analyze", "let's see", "let's look"
    ]

    for p in doc.paragraphs:
        t = p.text.strip()
        if p.style.name.startswith('Heading 2'): h2 += 1
        if "🚨 TRAP" in t: trap += 1
        if "💡 TRICK" in t: trick += 1
        if "📝 QUOTE" in t: quote += 1
        if "[⚡ Quick Rev]" in t: rev += 1
        if "Visual anchor" in t: vis_fail += 1
        # Skip banned attribution check if paragraph is styled as a quote or contains "📝 QUOTE"
        if "📝 QUOTE" in t or p.style.name.lower() in ("quote", "intense quote", "block quote"):
            continue

        for b in banned:
            # Use word boundary search to avoid matching substrings like "we analyze" in "we analyzed"
            pattern = r'\b' + re.escape(b.lower()) + r'\b'
            if re.search(pattern, t.lower()):
                attr_fail += 1
                logging.warning(f"[FAIL] Gate 1: Banned attribution: '{b}' in: '{t[:80]}...'")

    img_count = sum(1 for rel in doc.part.rels.values() if 'image' in rel.reltype)
    total_map_ex = sum(len(b.get('examples',[])) for b in concept_blocks)
    doc_ex = count_worked_examples_in_docx(doc)

    docx_metadata = build_docx_metadata(docx_path)

    # Calculate expected images. Instead of using a transient sidecar file (loophole)
    # or counting all extracted frames (including duplicates/logos), we compute it
    # directly by resolving the unique visual moments mapped in the concept block map.
    timestamp_to_frame = {f.get('timestamp'): name for name, f in frames.items() if f.get('timestamp')}
    unique_expected_imgs = set()
    for b in concept_blocks:
        for vm in b.get('visual_moments', []):
            ts = vm.get('timestamp', '').rstrip('*')
            v_type = vm.get('type', 'board')
            if v_type == 'slide':
                slide_num = vm.get('slide_number')
                if slide_num and slides:
                    is_ref = vm.get('source') == 'reference'
                    for s in slides:
                        if s.get('slide_number') == slide_num:
                            s_path = s.get('image_path', '')
                            if is_ref and 'reference_pages' not in s_path:
                                continue
                            if not is_ref and 'reference_pages' in s_path:
                                continue
                            unique_expected_imgs.add(os.path.basename(s_path))
                            break
            else:
                if ts and ts in timestamp_to_frame:
                    frame_fname = timestamp_to_frame[ts]
                    unique_expected_imgs.add(frame_fname)
                else:
                    # Fallback pattern matching
                    fallback_name = f"{b.get('block_id', '')}_{ts.replace(':', '')}.jpg"
                    if fallback_name in frames:
                        unique_expected_imgs.add(fallback_name)
                    elif f"{fallback_name}*" in frames:
                        unique_expected_imgs.add(f"{fallback_name}*")
                        
    has_visual_sources = False
    input_dir = "lecture-input"
    if "fixtures" in concept_map_path or "note_quality" in concept_map_path:
        input_dir = os.path.dirname(os.path.abspath(concept_map_path))

    for ext in ['.mp4', '.mkv', '.avi']:
        if glob.glob(os.path.join(input_dir, f"*{ext}")):
            has_visual_sources = True
            break
    if os.path.exists(os.path.join(input_dir, "SLIDES.pdf")) or os.path.exists(os.path.join(input_dir, "SLIDES.pptx")):
        has_visual_sources = True
    if os.path.exists(os.path.join(input_dir, "REFERENCE_NOTES.pdf")):
        has_visual_sources = True

    if has_visual_sources:
        exp_img = len(unique_expected_imgs)
    else:
        exp_img = 0

    undisc = [s for s in slides if not s.get('discussed',True)]
    all_text = get_all_text_from_docx(doc)

    # Gate 12: Exercise questions must contain real text, not just integers
    empty_exercise_count = 0
    for b in concept_blocks:
        exercises = b.get('exercise_questions', [])
        for eq in exercises:
            if isinstance(eq, int) or (isinstance(eq, str) and not eq.strip()):
                empty_exercise_count += 1
    if empty_exercise_count > 0:
        logging.warning(f"[FAIL] Gate 12: Found {empty_exercise_count} empty or invalid exercise questions.")

    # Gate 13: Quotes must not contain raw SRT artifacts (garbled text, mid-sentence starts)

    srt_artifact_pattern = re.compile(r'^\s*\d+\s+\d{2}:\d{2}:\d{2},\d{3}\s*-->')
    bad_quotes = 0
    for b in concept_blocks:
        for q in b.get('teacher_quotes', []):
            q_clean = q.strip()
            if srt_artifact_pattern.search(q_clean) or q_clean.startswith(('-->', ' ', 'ा', 'ि', 'े')):
                bad_quotes += 1
                logging.warning(f"[FAIL] Gate 13: Quote with SRT artifact: '{q_clean[:80]}...'")

    # Gate 14: Concept block titles must be meaningful (not just question ranges)
    meaningful_title_pattern = re.compile(r'^.*\(Questions?\s*\d+')
    bad_titles = 0
    for b in concept_blocks:
        if meaningful_title_pattern.match(b.get('title', '')):
            bad_titles += 1
            logging.warning(f"[FAIL] Gate 14: Generic title: '{b.get('title')}'")

    # Gate 15: Explanation Conciseness check
    verbose_explanations = 0
    for b in concept_blocks:
        expl = b.get('explanation', '')
        if len(expl) > 4000:
            verbose_explanations += 1
            logging.warning(f"[FAIL] Gate 15: Verbose explanation in {b.get('block_id', '?')}: {len(expl)} chars")
    
    # Gate 9: Slide Handling - fail if slide_manifest is empty when slides are expected
    # Also check for undiscussed slides with OCR text appearing in document
    slide_file_exists = any(os.path.exists(os.path.join(input_dir, f)) for f in ["slides.pdf", "SLIDES.pdf", "slides.PDF", "SLIDES.PDF"])
    has_slides_expected = len(concept_blocks) > 0 and slide_file_exists
    slide_manifest_empty = len(slides) == 0
    
    # Gate 9 fails if: (1) slide manifest is empty when concept blocks exist and slides are expected, OR (2) undiscussed slides found in doc
    undisc_slide_in_doc = any((ot := s.get('ocr_text', '').strip()) and len(ot) > 5 and ot in all_text for s in undisc)
    
    # Fixed: Gate 9 now FAILS if slide_manifest is empty when we expect slides (concept_blocks exist and slides are present)
    if has_slides_expected and slide_manifest_empty:
        logging.warning("[FAIL] Gate 9: Slide manifest is empty but concept blocks exist - missing slide data")
        gate_9_result = False
    else:
        if undisc_slide_in_doc:
            logging.warning("[FAIL] Gate 9: Undiscussed slide text appears in document.")
        gate_9_result = not undisc_slide_in_doc
    
    # Gate 2: Revision box is optional, but if present, must be <= h2 count
    gate_2_result = (rev <= h2) and (h2 > 0)
    
    # Gate 16: Table Presence check
    has_table_defined = any('table' in b for b in concept_blocks)
    gate_16_result = True
    if has_table_defined and len(doc.tables) == 0:
        logging.warning("[FAIL] Gate 16: Concept map contains a table definition but docx has 0 tables.")
        gate_16_result = False

    # Gate 17: Sequence Integrity check
    gate_17_result = True
    h2_texts = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith('Heading 2')]
    expected_h2_texts = [b.get('title', '').strip() for b in concept_blocks]
    for i, h2_text in enumerate(h2_texts):
        if i < len(expected_h2_texts) and expected_h2_texts[i] not in h2_text:
            logging.warning(f"[FAIL] Gate 17: Heading sequence mismatch at index {i}. Expected: '{expected_h2_texts[i]}', Got: '{h2_text}'")
            gate_17_result = False

    # Gate 18: Exact Worked Examples check
    # generate_docx.py renders scenario_or_problem (falling back to sentence),
    # so we check both fields. For multi-line SQL/code examples that get split
    # across paragraphs, use word-level coverage as fallback (≥80% significant words).
    gate_18_result = True
    missing_examples = []
    norm_doc = "".join(c.lower() for c in all_text if c.isalnum())
    all_text_lower = all_text.lower()

    def _norm(text):
        no_tags = re.sub(r'<[^>]+>', '', text)
        cleaned = clean_attributions(no_tags)
        formatted = format_math_text(cleaned)
        return "".join(c.lower() for c in formatted if c.isalnum())

    def _word_coverage(text, doc_text_lower):
        """Check if ≥80% of significant words (len≥3) from text appear in doc."""
        words = [w for w in re.findall(r'[a-zA-Z0-9]{3,}', text.lower())]
        if not words:
            return True
        matched = sum(1 for w in words if w in doc_text_lower)
        return (matched / len(words)) >= 0.80

    for b in concept_blocks:
        for ex in b.get('examples', []):
            found = False
            # Try both sentence and scenario_or_problem (what docx actually renders)
            candidates = []
            for field in ('sentence', 'scenario_or_problem'):
                val = ex.get(field, '').strip()
                if val:
                    candidates.append(val)

            for cand in candidates:
                norm_cand = _norm(cand)
                if not norm_cand or norm_cand in norm_doc:
                    found = True
                    break

            # Fallback: word-level coverage for multi-line SQL/code examples
            if not found:
                for cand in candidates:
                    if _word_coverage(cand, all_text_lower):
                        found = True
                        break

            if not found:
                missing_examples.append(candidates[0] if candidates else '(empty)')
    if missing_examples:
        logging.warning(f"[FAIL] Gate 18: Missing exact worked examples in document: {missing_examples}")
        gate_18_result = False

    # Calculate Friction Index
    fi, tot_w, clozes, cues = calculate_friction_index(doc, concept_blocks)
    gate_friction_ok = fi <= 0.40
    if not gate_friction_ok:
        logging.warning(f"[FAIL] Gate 19: Friction Index {fi:.3f} exceeds maximum allowed of 0.40")

    doc_title = next((p.text.strip() for p in doc.paragraphs if p.text.strip()), "")
    def norm_title(value):
        return re.sub(r'[^a-z0-9]+', '', value.lower())
    source_trace_ok = True
    if concept_map_title and norm_title(concept_map_title) != norm_title(doc_title):
        logging.warning(f"[FAIL] Gate 8: Document title '{doc_title}' does not match concept map title '{concept_map_title}'")
        source_trace_ok = False

    # Gate 20: Transcript Coverage
    gate_20_result = True
    transcript_path = os.path.join(input_dir, "transcript.srt")
    if not os.path.exists(transcript_path):
        for name in ["transcript.txt", "transcript.vtt", "TRANSCRIPT.srt", "TRANSCRIPT.txt", "TRANSCRIPT.vtt"]:
            p_path = os.path.join(input_dir, name)
            if os.path.exists(p_path):
                transcript_path = p_path
                break
            
    if os.path.exists(transcript_path):
        try:
            # Calculate range coverage percent using mathematical union of intervals
            covered_percentages = set()
            for block in concept_blocks:
                rng = block.get('transcript_range_percent', [0, 0])
                if len(rng) == 2:
                    for p_val in range(int(rng[0]), int(rng[1])):
                        covered_percentages.add(p_val)
            total_coverage_pct = len(covered_percentages)
            
            # Check heading presence in docx
            h2_texts = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith('Heading 2')]
            expected_h2_texts = [b.get('title', '').strip() for b in concept_blocks]
            found_headings_count = sum(1 for exp in expected_h2_texts if any(exp in h2 for h2 in h2_texts))
            
            heading_coverage_pct = (found_headings_count / len(concept_blocks) * 100) if concept_blocks else 100
            
            logging.info(f"Gate 20: Transcript coverage from concept map = {total_coverage_pct}%. Heading presence in docx = {found_headings_count}/{len(concept_blocks)} ({heading_coverage_pct:.1f}%).")
            
            if total_coverage_pct < 80:
                logging.warning(f"[FAIL] Gate 20: Transcript coverage in concept map is only {total_coverage_pct}% (needs >= 80%).")
                gate_20_result = False
            if heading_coverage_pct < 80:
                logging.warning(f"[FAIL] Gate 20: Heading coverage in docx is only {heading_coverage_pct:.1f}% (needs >= 80%).")
                gate_20_result = False
        except Exception as e:
            logging.warning(f"Error executing Gate 20 audit check: {e}")
            gate_20_result = False
    else:
        logging.warning("Gate 20: Transcript file not found. Skipping coverage check.")

    # Gate 21: Hindi/Hinglish Constraint
    gate_21_result = True
    
    # Check for Devanagari script
    devanagari_characters = re.findall(r'[\u0900-\u097F]', all_text)
    if len(devanagari_characters) > 0:
        logging.warning(f"[FAIL] Gate 21: Devanagari script detected in notes ({len(devanagari_characters)} characters). Hindi script is forbidden.")
        gate_21_result = False
        
    # Check for transliterated Hinglish keywords
    hindi_keywords = {'hai', 'ki', 'aur', 'toh', 'bhi', 'mein', 'se', 'ka', 'ko', 'ye', 'wo', 'kya', 'tha', 'thi', 'karna', 'karo', 'liye', 'nahi', 'ab', 'jab', 'tab'}
    hindi_count = 0
    # Use raw all_text to preserve spaces for discrete word boundaries, not the squashed norm_doc
    words = re.findall(r'\b[a-z]+\b', all_text.lower())
    for word in words:
        if word in hindi_keywords:
            hindi_count += 1
            
    if hindi_count > 40:
        logging.warning(f"[FAIL] Gate 21: Excessive Hindi/Hinglish detected in document ({hindi_count} common Hindi words found). Max allowed is 40.")
        gate_21_result = False

    # Gate 22: Styling and Highlighting Conformity
    gate_22_result = True
    approved_pastels = {
        'FFF2CC', 'E8F8F5', 'E1F5FE', 'F1F5F9', 'FEE2E2', 'FFEDD5', 'F3E8FF',
        'AED6F1', 'F5B7B1', 'D2B4DE', 'A3E4D7', 'F5CBA7', 'D5D8DC', 'A9CCE3',
        'EAECEE'
    }
    
    all_paras = get_all_paragraphs(doc)
    for p in all_paras:
        t = p.text.strip()
        pPr = p._p.pPr
        shd = None if pPr is None else pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
        
        if t.startswith("[⚡ Quick Rev]"):
            if shd is None:
                logging.warning("[FAIL] Gate 22: Quick Revision box has no shading.")
                gate_22_result = False
            else:
                fill_val = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                if not fill_val or fill_val.upper().replace('#', '') != 'D6EAF8':
                    logging.warning(f"[FAIL] Gate 22: Quick Revision box has invalid shading fill: '{fill_val}' (expected 'D6EAF8').")
                    gate_22_result = False
                    
        elif t.startswith("💡 Student Note / Doubt") or t.startswith("📝 Student Note:"):
            if is_paragraph_in_table(p):
                continue
            if shd is None:
                logging.warning(f"[FAIL] Gate 22: Student Note box starting with '{t[:20]}...' has no shading.")
                gate_22_result = False
            else:
                fill_val = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                if not fill_val or fill_val.upper().replace('#', '') != 'F0F4F8':
                    logging.warning(f"[FAIL] Gate 22: Student Note box starting with '{t[:20]}...' has invalid shading fill: '{fill_val}' (expected 'F0F4F8').")
                    gate_22_result = False
                    
        for r in p.runs:
            rPr = r._r.rPr
            highlight = None if rPr is None else rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}highlight')
            if highlight is not None:
                val = highlight.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'unknown')
                logging.warning(f"[FAIL] Gate 22: Native Word highlight '{val}' found on run '{r.text[:30]}...'. Native Word highlights are strictly banned.")
                gate_22_result = False
                
            shd_run = None if rPr is None else rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
            if shd_run is not None:
                fill_val = shd_run.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                if not fill_val or fill_val.upper().replace('#', '') not in approved_pastels:
                    logging.warning(f"[FAIL] Gate 22: Run '{r.text[:30]}...' has unapproved run shading: '{fill_val}'.")
                    gate_22_result = False

    if h2 == 0:
        logging.warning("[FAIL] Gate 1: Missing structural elements (h2=0).")
    if rev > h2:
        logging.warning(f"[FAIL] Gate 1: Structural mismatch, revision boxes ({rev}) > topics ({h2}).")
    if vis_fail > 0:
        logging.warning(f"[FAIL] Gate 1: Detected {vis_fail} instances of raw visual anchor text.")

    if not gate_2_result:
        logging.warning("[FAIL] Gate 2: Revision box is not properly placed relative to headings.")
        
    if concept_blocks and h2 != len(concept_blocks):
        logging.warning(f"[FAIL] Gate 3: Chronological Flow mismatch. Expected {len(concept_blocks)} headings, found {h2}.")
        
    if len(concept_blocks) == 0:
        logging.warning("[FAIL] Gate 4: Content Completeness - No concept blocks provided.")
        
    if doc_ex < total_map_ex:
        logging.warning(f"[FAIL] Gate 5: Factual Accuracy - Example count mismatch. Expected {total_map_ex}, found {doc_ex}.")
        
    if vis_fail > 0:
        logging.warning(f"[FAIL] Gate 6: Image Integrity - Detected {vis_fail} visual anchors without proper rendering.")
        
    if h2 < 1 or img_count < exp_img * 0.8:
        logging.warning(f"[FAIL] Gate 7: Minimum Counts - Expected >= 1 heading and >={int(exp_img * 0.8)} images, got {h2} headings and {img_count} images.")
        
    if total_map_ex and doc_ex < total_map_ex:
        logging.warning(f"[FAIL] Gate 10: Example Coverage - Extracted {doc_ex}/{total_map_ex} examples.")
        
    if exp_img and img_count < exp_img * 0.8:
        logging.warning(f"[FAIL] Gate 11: Visual Coverage - Inserted {img_count}/{exp_img} expected images.")

    gates = {
        'Gate 1: Structural Integrity':    h2 > 0 and rev <= h2 and vis_fail == 0 and attr_fail == 0,
        'Gate 2: Revision Box Placement':  gate_2_result,
        'Gate 3: Chronological Flow':      h2 == len(concept_blocks) if concept_blocks else False,
        'Gate 4: Content Completeness':    len(concept_blocks) > 0,
        'Gate 5: Factual Accuracy':        doc_ex >= total_map_ex,
        'Gate 6: Image Integrity':         vis_fail == 0,
        'Gate 7: Minimum Counts':          h2 >= 1 and img_count >= exp_img * 0.8,
        'Gate 8: Source Traceability':     source_trace_ok,
        'Gate 9: Slide Handling':          gate_9_result,
        'Gate 10: Example Coverage':       doc_ex >= total_map_ex if total_map_ex else True,
        'Gate 11: Visual Coverage':        img_count >= exp_img * 0.8 if exp_img else True,
        'Gate 12: Exercise Content':       empty_exercise_count == 0,
        'Gate 13: Quote Quality':          bad_quotes == 0,
        'Gate 14: Meaningful Titles':      bad_titles == 0,
        'Gate 15: Explanation Conciseness': verbose_explanations == 0,
        'Gate 16: Table Presence':          gate_16_result,
        'Gate 17: Sequence Integrity':      gate_17_result,
        'Gate 18: Exact Worked Examples':  gate_18_result,
        'Gate 19: Friction Index Constraint': gate_friction_ok,
        'Gate 20: Transcript Coverage':     gate_20_result,
        'Gate 21: English Enforcement':     gate_21_result,
        'Gate 22: Styling and Highlighting Conformity': gate_22_result,
    }

    # Gate 23: Word Count Budget (Strict 4,000-word ceiling for standard lectures)
    max_budget = 4000
    gates['Gate 23: Word Count Budget'] = tot_w <= max_budget
    if tot_w > max_budget:
        logging.warning(f"[FAIL] Gate 23: Word count is {tot_w} (max: {max_budget}). Notes are too verbose.")
    elif tot_w > (max_budget * 0.8):
        logging.warning(f"[WARN] Gate 23: Word count is {tot_w} (target: 2500-3500). Slightly verbose but passing.")
    else:
        logging.info(f"[INFO] Gate 23: Word count is {tot_w} — within target student note range.")

    # Gate 24: Total Callout Box Cap (strict limit of 6)
    total_callouts = trap + trick + quote
    gates['Gate 24: Callout Box Cap'] = total_callouts <= 6
    if total_callouts > 6:
        logging.warning(f"[FAIL] Gate 24: Total callout boxes is {total_callouts} (max: 6). Too many callout boxes.")
    else:
        logging.info(f"[INFO] Gate 24: Total callout boxes is {total_callouts} — within limits.")

    # Gate 25: Short Note Audit
    short_note_ok = True
    short_note_path = docx_path.replace(".docx", "_SHORTNOTE.md")
    if not os.path.exists(short_note_path) and "LECTURE_NOTES.docx" in docx_path:
        alt_path = os.path.join(os.path.dirname(docx_path), "LECTURE_SHORTNOTE.md")
        if os.path.exists(alt_path):
            short_note_path = alt_path

    if os.path.exists(short_note_path):
        try:
            with open(short_note_path, "r", encoding="utf-8") as f:
                sn_content = f.read()
            
            # 1. No Devanagari script
            if re.search(r'[\u0900-\u097F]', sn_content):
                logging.warning("[FAIL] Gate 25: Devanagari script detected in short note. Hindi script is forbidden.")
                short_note_ok = False
                
            # 2. Context anchor check
            if "answering:" not in sn_content or "From " not in sn_content:
                logging.warning("[FAIL] Gate 25: Short note lacks a valid context anchor (From **[Lecture Title]**, answering: ...)")
                short_note_ok = False
                
            # 3. Answer leakage check in Self-Test
            if "self-test" in sn_content.lower():
                self_test_part = sn_content.lower().split("self-test")[1].strip()
                self_test_part = self_test_part.lstrip(":-#* \t\n")
                if "source" in self_test_part:
                    self_test_part = self_test_part.split("source")[0].strip()
                if "##" in self_test_part:
                    self_test_part = self_test_part.split("##")[0].strip()
                if len(self_test_part) > 600 or "|" in self_test_part:
                    logging.warning("[FAIL] Gate 25: Self-Test section in short note contains leaked answers or excessive formatting.")
                    short_note_ok = False
            else:
                logging.warning("[FAIL] Gate 25: Self-Test section missing in short note.")
                short_note_ok = False
                
            # 4. Word count check (between 50 and 600 words)
            sn_words = len(re.sub(r'[^\w\s]', '', sn_content).split())
            if sn_words < 50 or sn_words > 600:
                logging.warning(f"[FAIL] Gate 25: Short note word count is {sn_words} (expected 50-600 words).")
                short_note_ok = False
            else:
                logging.info(f"Gate 25: Short note word count is {sn_words} (within limits).")
                
        except Exception as e:
            logging.warning(f"[FAIL] Gate 25: Error reading short note file: {e}")
            short_note_ok = False
    else:
        if "note_quality" in docx_path or "fixtures" in docx_path:
            logging.warning(f"[FAIL] Gate 25: Short note file not found at {short_note_path}.")
            short_note_ok = False
        else:
            logging.info("Gate 25: Short note file not found (disabled or skipped). Passing Gate 25 by default.")
            short_note_ok = True
        
    gates['Gate 25: Short Note Audit'] = short_note_ok

    logging.info("\n--- AUDIT RESULTS ---")
    logging.info(f"H2: {h2}  RevBox: {rev}  Traps: {trap}  Tricks: {trick}  Quotes: {quote}  Tables: {len(doc.tables)}")
    logging.info(f"Images: {img_count}  MapEx: {total_map_ex}  DocEx: {doc_ex}  AttrFail: {attr_fail}")
    logging.info(f"Friction Index: {fi:.3f} (Words: {tot_w}, Clozes: {clozes}, Cues: {cues})")
    logging.info(f"Word Count: {tot_w} (Target: 2,500-3,500)")
    logging.info("---------------------")
    all_ok = True
    for g, ok in gates.items():
        logging.info(f"{'[PASS]' if ok else '[FAIL]'} {g}")
        if not ok:
            all_ok = False
    
    if all_ok:
        logging.info("\n[SUCCESS] All gates passed.")
    else:
        logging.warning("\n[FAIL] Some gates failed.")
        
    return all_ok, gates, audit_feedback

if __name__ == '__main__':
    p = argparse.ArgumentParser()
    p.add_argument('--docx', default='notes-output/LECTURE_NOTES.docx')
    p.add_argument('--concept-map', default='concept_block_map.json')
    p.add_argument('--frame-manifest', default='frame_manifest.json')
    p.add_argument('--slide-manifest', default='slide_manifest.json')
    args = p.parse_args()
    success, gates, feedback_data = run_audit(args.docx, args.concept_map, args.frame_manifest, args.slide_manifest)
    numeric_gates = {}
    for gate_name, ok in gates.items():
        match = __import__('re').search(r'Gate\s+(\d+):', gate_name)
        if match:
            numeric_gates[match.group(1)] = ok
    try:
        payload = {
            "_docx": build_docx_metadata(args.docx),
            "_timestamp": __import__("datetime").datetime.now().isoformat(),
            **numeric_gates,
        }
        with open("logs/last_run_audit.json", "w", encoding="utf-8") as f:
            json.dump(payload, f, indent=2)
            
        with open("logs/audit_feedback.json", "w", encoding="utf-8") as f:
            json.dump(feedback_data, f, indent=2)
    except Exception as e:
        logging.warning(f"Could not write logs: {e}")
    sys.exit(0 if success else 1)
