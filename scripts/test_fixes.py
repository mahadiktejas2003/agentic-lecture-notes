#!/usr/bin/env python3
import unittest
import re
from docx import Document
from docx.shared import Pt, RGBColor
from scripts.generate_docx import (
    add_rich_runs,
    format_math_text,
    add_revision_box,
    add_cornell_block,
    add_styled_table,
    add_shaded_formula,
    add_formatted_explanation_paragraphs
)

class TestDocxFixes(unittest.TestCase):

    def test_r1_revision_box_color_forcing(self):
        doc = Document()
        # Normal bullet, custom color cloze bullet, and standard bullet
        bullets = ["Normal text here", "Text with <cloze>cloze word</cloze> here", "More text"]
        p = add_revision_box(doc, bullets, rule="X + Y = Z")
        
        # Verify runs
        runs = p.runs
        self.assertTrue(len(runs) > 0)
        for run in runs:
            if run.text.strip().isdigit() and run.font.superscript:
                self.assertEqual(run.font.color.rgb, RGBColor(0x00, 0x70, 0xC0))
            elif "cloze word" in run.text:
                self.assertEqual(run.font.color.rgb, RGBColor(0x00, 0x70, 0xC0))
            else:
                self.assertEqual(run.font.color.rgb, RGBColor(0, 0, 0))

    def test_r1_cornell_block_color_forcing(self):
        doc = Document()
        content = [
            "💡 Student Note / Doubt: This is a test doubt.",
            "💡 Teacher's Intuition & Analogies: This is intuition.",
            "Normal text without shading"
        ]
        add_cornell_block(doc, "Cue Text", content)
        
        # Find the table and cells
        table = doc.tables[-1]
        cell_note = table.cell(0, 1)
        
        # We rendered 3 paragraphs for content in cell_note
        paragraphs = cell_note.paragraphs
        self.assertEqual(len(paragraphs), 3)
        
        # First paragraph (Student Note / Doubt) -> shading F0F4F8, runs are black
        p1 = paragraphs[0]
        for r in p1.runs:
            self.assertEqual(r.font.color.rgb, RGBColor(0, 0, 0))
            
        # Second paragraph (Teacher's Intuition) -> shading FFF2CC, runs are black
        p2 = paragraphs[1]
        for r in p2.runs:
            self.assertEqual(r.font.color.rgb, RGBColor(0, 0, 0))

    def test_r1_styled_table_alternate_row_color(self):
        doc = Document()
        headers = ["H1", "H2"]
        rows = [
            ["Row0 C0", "Row0 C1"],
            ["Row1 C0", "Row1 C1 <cloze>cloze</cloze>"], # Alternate row (shaded)
            ["Row2 C0", "Row2 C1"]
        ]
        add_styled_table(doc, headers, rows)
        
        table = doc.tables[-1]
        row_shaded = table.rows[2]
        
        p0 = row_shaded.cells[0].paragraphs[0]
        for run in p0.runs:
            self.assertEqual(run.font.color.rgb, RGBColor(0, 0, 0))
            
        p1 = row_shaded.cells[1].paragraphs[0]
        for run in p1.runs:
            if run.text.strip().isdigit() and run.font.superscript:
                self.assertEqual(run.font.color.rgb, RGBColor(0x00, 0x70, 0xC0))
            elif "cloze" in run.text:
                self.assertEqual(run.font.color.rgb, RGBColor(0x00, 0x70, 0xC0))
            else:
                self.assertEqual(run.font.color.rgb, RGBColor(0, 0, 0))

    def test_r1_shaded_formula(self):
        doc = Document()
        formula = "E = mc^2"
        p = add_shaded_formula(doc, formula)
        
        # Check center alignment
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        self.assertEqual(p.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        
        # Check runs are black
        for run in p.runs:
            self.assertEqual(run.font.color.rgb, RGBColor(0, 0, 0))

    def test_r2_latex_subscripts_adjacent_text(self):
        doc = Document()
        p = doc.add_paragraph()
        
        # Adjacent subscripts tests
        cases = [
            (r"d\text{min}", "d<sub>min</sub>"),
            (r"d_\text{min}", "d<sub>min</sub>"),
            (r"d_{\text{min}}", "d<sub>min</sub>")
        ]
        for inp, expected in cases:
            p.text = ""
            add_rich_runs(p, inp)
            runs_repr = "".join(f"<sub>{r.text}</sub>" if r.font.subscript else r.text for r in p.runs)
            self.assertEqual(runs_repr, expected)

    def test_r2_robust_math_operators(self):
        doc = Document()
        
        # Verify replacement to Unicode equivalents
        operators = [
            (r"\min", "min"),
            (r"\pm", "±"),
            (r"\times", "×"),
            (r"\leq", "≤"),
            (r"\geq", "≥"),
            (r"\neq", "≠"),
            (r"\approx", "≈"),
            (r"\div", "÷"),
            (r"\cdot", "·"),
            (r"\rightarrow", "→"),
            (r"\leftarrow", "←"),
            (r"\Rightarrow", "⇒"),
            (r"\infty", "∞"),
            (r"\max", "max")
        ]
        for op, expected in operators:
            p = doc.add_paragraph()
            add_rich_runs(p, op)
            self.assertEqual(p.runs[0].text, expected)

        # Verify unbraced and braced subscripts participation (e.g. x_\min -> x_min -> x<sub>min</sub>)
        p = doc.add_paragraph()
        add_rich_runs(p, r"x_\min")
        runs_repr = "".join(f"<sub>{r.text}</sub>" if r.font.subscript else r.text for r in p.runs)
        self.assertEqual(runs_repr, "x<sub>min</sub>")

    def test_r2_greek_letter_backslash(self):
        # format_math_text should convert both alpha and \alpha to α
        self.assertEqual(format_math_text("alpha"), "α")
        self.assertEqual(format_math_text(r"\alpha"), "α")
        self.assertEqual(format_math_text("beta"), "β")
        self.assertEqual(format_math_text(r"\beta"), "β")

    def test_r3_paragraph_splitting(self):
        doc = Document()
        # A line longer than 350 characters
        long_line = (
            "This is sentence one which is designed to be quite long so that the total character count of the paragraph exceeds the threshold of 350 characters. "
            "This is sentence two which also contains a lot of explanation text and details about the mathematical model we are discussing. "
            "This is sentence three, providing further intuitive examples and explanations to ensure the student understands the concepts. "
            "This is sentence four, detailing the potential traps and edge cases in the calculations. "
            "This is sentence five, showing the step-by-step algebra derivation. "
            "This is sentence six, which completes this long explanation paragraph."
        )
        self.assertTrue(len(long_line) > 350)
        
        # Run paragraph splitting
        add_formatted_explanation_paragraphs(doc, long_line)
        
        # Sentences are grouped into chunks of at most 3 sentences:
        # Paragraph 1: "This is sentence one which is designed to be quite long so that the total character count of the paragraph exceeds the threshold of 350 characters. This is sentence two which also contains a lot of explanation text and details about the mathematical model we are discussing. This is sentence three, providing further intuitive examples and explanations to ensure the student understands the concepts."
        # Paragraph 2: "This is sentence four, detailing the potential traps and edge cases in the calculations. This is sentence five, showing the step-by-step algebra derivation. This is sentence six, which completes this long explanation paragraph."
        self.assertEqual(len(doc.paragraphs), 2)
        self.assertEqual(
            doc.paragraphs[0].text, 
            "This is sentence one which is designed to be quite long so that the total character count of the paragraph exceeds the threshold of 350 characters. This is sentence two which also contains a lot of explanation text and details about the mathematical model we are discussing. This is sentence three, providing further intuitive examples and explanations to ensure the student understands the concepts."
        )
        self.assertEqual(
            doc.paragraphs[1].text,
            "This is sentence four, detailing the potential traps and edge cases in the calculations. This is sentence five, showing the step-by-step algebra derivation. This is sentence six, which completes this long explanation paragraph."
        )

if __name__ == "__main__":
    unittest.main()
