import os
import sys
import subprocess
import docx
import json

def run_cmd(cmd):
    res = subprocess.run(cmd, shell=True, capture_output=True, text=True)
    return res

def main():
    print("=== Running Comprehensive Note Quality Diagnostic Tests ===")
    
    # Paths
    concept_map = "tests/fixtures/note_quality/technical_fixture.json"
    frame_manifest = "tests/fixtures/note_quality/frame_manifest.json"
    slide_manifest = "tests/fixtures/note_quality/slide_manifest.json"
    out_docx = "tests/fixtures/note_quality/notes.docx"
    out_short_note = "tests/fixtures/note_quality/notes_SHORTNOTE.md"
    
    # Clean previous outputs
    for path in [out_docx, out_short_note]:
        if os.path.exists(path):
            os.remove(path)
            
    # -------------------------------------------------------------
    # Test Case 1: Standard Generation & Validation
    # -------------------------------------------------------------
    print("\n[Test 1] Generating and Auditing standard fixture...")
    cmd_docx = f"venv/bin/python scripts/generate_docx.py --concept-map {concept_map} --frame-manifest {frame_manifest} --slide-manifest {slide_manifest} --output {out_docx}"
    res = run_cmd(cmd_docx)
    assert res.returncode == 0, f"generate_docx.py failed: {res.stderr}"
    
    cmd_short = f"venv/bin/python scripts/generate_short_note.py --concept-map {concept_map} --slide-manifest {slide_manifest} --frame-manifest {frame_manifest} --mode heuristic --output {out_short_note}"
    res = run_cmd(cmd_short)
    assert res.returncode == 0, f"generate_short_note.py failed: {res.stderr}"
    
    # Run audit on correct outputs
    cmd_audit = f"venv/bin/python scripts/audit.py --docx {out_docx} --concept-map {concept_map} --frame-manifest {frame_manifest} --slide-manifest {slide_manifest}"
    res_audit = run_cmd(cmd_audit)
    assert res_audit.returncode == 0, f"Standard audit failed to pass: {res_audit.stdout}\n{res_audit.stderr}"
    
    # -------------------------------------------------------------
    # Test Case 2: Inspect DOCX for formatting remediation rules
    # -------------------------------------------------------------
    print("[Test 2] Inspecting DOCX formatting details...")
    doc = docx.Document(out_docx)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    # Assert Section 1 & 2 headers are absent
    has_flow_outline = any("Section 1: Lecture Flow Outline" in p for p in paragraphs)
    has_detailed_blocks = any("Section 2: Detailed Concept Blocks" in p for p in paragraphs)
    assert not has_flow_outline, "Flow outline scaffolding still present in DOCX"
    assert not has_detailed_blocks, "Detailed blocks section header still present in DOCX"
    
    # Assert H2 headings do not contain "CBx:" prefix
    h2_headings = [p.text for p in doc.paragraphs if p.style.name == 'Heading 2']
    for h2 in h2_headings:
        assert not h2.startswith("CB"), f"Heading 2 should not have CBx prefix: {h2}"
        
    # Assert no visual moment fallback captions
    for p in paragraphs:
        assert "[Visual Moment at" not in p, f"Found positional visual fallback in text: {p}"
        
    # -------------------------------------------------------------
    # Test Case 3: Verify Gate 25 - Absent Short Note Failure (RED)
    # -------------------------------------------------------------
    print("[Test 3] Verifying absent short-note failure (Gate 25 Red)...")
    if os.path.exists(out_short_note):
        os.remove(out_short_note)
        
    res_absent = run_cmd(cmd_audit)
    assert res_absent.returncode != 0, "Audit passed when short-note file was absent!"
    assert "Gate 25: Short note file not found" in res_absent.stderr or "Gate 25: Short note file not found" in res_absent.stdout, "Audit did not log correct Gate 25 failure message for missing file."

    # -------------------------------------------------------------
    # Test Case 4: Verify Gate 25 - Answer Leakage Failure (RED)
    # -------------------------------------------------------------
    print("[Test 4] Verifying answer leakage detection (Gate 25 Red)...")
    leak_content = """# Title
## Context Anchor
From **Title**, answering: *test*

## Technical Core Concepts
- Concept: definition

## Active Recall Self-Test
1. Question 1?
| leaked answer | in table |

## Source Provenance
*Source: test*
"""
    with open(out_short_note, "w", encoding="utf-8") as f:
        f.write(leak_content)
        
    res_leak = run_cmd(cmd_audit)
    assert res_leak.returncode != 0, "Audit passed when short-note contained answer leakage!"
    assert "Gate 25: Self-Test section in short note contains leaked answers" in res_leak.stderr or "Gate 25: Self-Test section in short note contains leaked answers" in res_leak.stdout, "Audit did not log correct Gate 25 failure message for answer leakage."

    # -------------------------------------------------------------
    # Test Case 5: Verify Gate 25 - Word Count Ceiling (RED)
    # -------------------------------------------------------------
    print("[Test 5] Verifying short-note word count ceiling (Gate 25 Red)...")
    too_many_words = "word " * 650
    wordy_content = f"""# Title
## Context Anchor
From **Title**, answering: *test*

## Technical Core Concepts
- Concept: definition

## Active Recall Self-Test
1. Question 1?

## Source Provenance
*Source: test*

{too_many_words}
"""
    with open(out_short_note, "w", encoding="utf-8") as f:
        f.write(wordy_content)
        
    res_wordy = run_cmd(cmd_audit)
    assert res_wordy.returncode != 0, "Audit passed when short-note exceeded word limits!"
    assert "Gate 25: Short note word count is" in res_wordy.stderr or "Gate 25: Short note word count is" in res_wordy.stdout, "Audit did not log correct Gate 25 failure message for excessive word count."

    # -------------------------------------------------------------
    # Test Case 6: Empty Traps Section Handling
    # -------------------------------------------------------------
    print("[Test 6] Verifying empty traps section handling...")
    with open(concept_map, "r", encoding="utf-8") as f:
        cmap_data = json.load(f)
    cmap_data["blocks"][0]["traps"] = []
    
    temp_map = "tests/fixtures/note_quality/technical_fixture_no_traps.json"
    with open(temp_map, "w", encoding="utf-8") as f:
        json.dump(cmap_data, f)
        
    cmd_short_no_traps = f"venv/bin/python scripts/generate_short_note.py --concept-map {temp_map} --slide-manifest {slide_manifest} --frame-manifest {frame_manifest} --mode heuristic --output {out_short_note}"
    run_cmd(cmd_short_no_traps)
    
    with open(out_short_note, "r", encoding="utf-8") as f:
        no_traps_content = f.read()
        
    if os.path.exists(temp_map):
        os.remove(temp_map)
        
    assert "Trap / Mistake Log" not in no_traps_content, "Generated empty/fabricated trap section when traps list was empty."

    print("\n✅ All red-capable note-quality gate verification tests PASSED.")
    sys.exit(0)

if __name__ == "__main__":
    main()
